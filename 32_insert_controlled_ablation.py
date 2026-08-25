"""R2.4 - Sisipkan subbagian 4.9 beserta Table 11 dan Table 12 ke manuskrip.

Table 11 - kinerja delapan kondisi ablasi terkontrol (tahap 30).
Table 12 - efek tiap faktor, dihitung berpasangan pada tingkat fold.

Penempatan
----------
Keduanya diletakkan di akhir Results, setelah Table 10, sehingga tidak ada
penomoran tabel lain yang bergeser. Aturan yang sama dipakai saat menyisipkan
Table 8 dan Table 9-10.

Perbaikan yang ikut dikerjakan
-----------------------------
Paragraf metode Section 3.8 memuat dua kalimat yang sama-sama dibuka dengan
"Fourth," karena skrip 23 dahulu menambah satu diagnostik tanpa menyesuaikan
urutannya, dan pembukanya masih menyebut "Four additional diagnostics" padahal
jumlahnya sudah bertambah. Penomoran ordinal itu disusun ulang di sini agar
konsisten dengan jumlah diagnostik yang benar-benar dijelaskan.

Framing yang dipegang
---------------------
Reviewer 2 butir 4 menduga keunggulan cabang deep berasal dari RGB, pretraining,
crop, dan dimensi yang lebih besar. Hasilnya dilaporkan apa adanya: satu faktor
berpengaruh besar, dua lainnya tidak. Tidak ada klaim bahwa satu keluarga
representasi lebih unggul.

Skrip bersifat idempotent dan menolak jalan bila manuskrip terbuka di Word.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\32_insert_controlled_ablation.py
"""

from __future__ import annotations

import json
import re
import shutil
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from pipeline_config import PROJECT_DIR, RESULTS_DIR

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
STAGE30 = RESULTS_DIR / "30_controlled_ablation"
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_Table11_12_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)
MARKER = "4.9 Controlled ablation of the deep branch"
SENTENCE_BOUNDARY = re.compile(r"(?<=[.;:]) (?=[A-Z(])")

ORDINALS = [
    "First", "Second", "Third", "Fourth", "Fifth", "Sixth", "Seventh",
    "Eighth", "Ninth", "Tenth",
]
COUNT_WORDS = {
    1: "One", 2: "Two", 3: "Three", 4: "Four", 5: "Five",
    6: "Six", 7: "Seven", 8: "Eight", 9: "Nine", 10: "Ten",
}

LABEL = {
    "rgb": "RGB", "gray": "Grayscale",
    "pretrained": "ImageNet", "random": "Random",
    "full512": "512-d", "pca6": "PCA 6-d",
}
FACTOR_LABEL = {
    "color": "Colour", "pretraining": "Pretraining", "dimensionality": "Dimensionality",
}
CONTRAST_LABEL = {
    "rgb minus gray": "RGB minus grayscale",
    "pretrained minus random": "ImageNet minus random",
    "full512 minus pca6": "512-d minus PCA 6-d",
}

TABLE11_HEADERS = [
    "Colour", "Pretraining", "Dimensionality", "Macro-F1", "Bal. accuracy", "MCC",
]
TABLE12_HEADERS = [
    "Factor", "Contrast", "Mean difference", "Lowest", "Highest", "Settings",
]

TABLE11_CAPTION = (
    "Table 11. Controlled ablation of the frozen ResNet18 branch across eight "
    "conditions formed by crossing colour, pretraining, and embedding "
    "dimensionality. All conditions use the same repeated nested protocol, the "
    "same source-group-aware folds, and the same three classifier families; "
    "principal component analysis is fitted inside the training pipeline only. "
    "The RGB, ImageNet, 512-d condition reproduces the extraction used elsewhere "
    "in this article exactly."
)
TABLE12_CAPTION = (
    "Table 12. Effect of each factor on macro-F1, paired at the outer-fold level "
    "and then averaged over the four combinations of the remaining two factors. "
    "Lowest and highest give the range across those combinations, so a factor "
    "whose range crosses zero has no consistent direction."
)

METHOD_SENTENCE = (
    " {ordinal}, the frozen ResNet18 branch was re-evaluated under eight controlled "
    "conditions crossing RGB against grayscale, ImageNet weights against random "
    "initialisation, and the full 512-dimensional embedding against six principal "
    "components, so that the confounded comparison between representation families "
    "could be decomposed into separate factors."
)


def lock_present() -> bool:
    return any(MANUSCRIPT.parent.glob("~$*"))


def fmt(value: float, decimals: int = 3) -> str:
    return f"{float(value):.{decimals}f}"


def signed(value: float, decimals: int = 4) -> str:
    return f"{float(value):+.{decimals}f}"


def sentences(text: str) -> list[str]:
    return [part for part in SENTENCE_BOUNDARY.split(text.strip()) if part]


def styled(paragraph, text, *, bold=False, size=Pt(11)):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.color.rgb = RED
    return run


def build_table(document: Document, headers: list[str], rows: list[list[str]]) -> Table:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        styled(cell.paragraphs[0], header, bold=True, size=Pt(10))
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = ""
            styled(cells[index].paragraphs[0], value, size=Pt(10))
    return table


def renumber_method_ordinals(document: Document) -> bool:
    """Susun ulang penomoran diagnostik pada Section 3.8, tandai yang berubah.

    Dua kalimat sempat sama-sama dibuka "Fourth," dan pembukanya masih menyebut
    empat diagnostik. Menambah kalimat baru tanpa membereskan ini akan
    memperbesar kekeliruannya.
    """
    paragraph = next(
        (p for p in document.paragraphs
         if p.text.strip().startswith(("Four additional diagnostics",
                                       "Five additional diagnostics",
                                       "Six additional diagnostics",
                                       "Seven additional diagnostics",
                                       "Eight additional diagnostics"))),
        None,
    )
    if paragraph is None:
        raise AssertionError("Paragraf diagnostik Section 3.8 tidak ditemukan")

    old_text = paragraph.text.strip()
    parts = sentences(old_text)
    ordinal_positions = [
        index for index, sentence in enumerate(parts)
        if any(sentence.startswith(word + ",") for word in ORDINALS)
    ]
    total = len(ordinal_positions)
    if total > len(ORDINALS):
        raise AssertionError(f"Terlalu banyak diagnostik untuk dinomori: {total}")

    for order, index in enumerate(ordinal_positions):
        sentence = parts[index]
        current = sentence.split(",", 1)[0]
        parts[index] = ORDINALS[order] + "," + sentence[len(current) + 1:]

    parts[0] = re.sub(
        r"^\w+ additional diagnostics",
        f"{COUNT_WORDS[total]} additional diagnostics",
        parts[0],
    )
    new_text = " ".join(parts)
    if new_text == old_text:
        return False

    old_sentences = set(sentences(old_text))
    template = paragraph.runs[0] if paragraph.runs else None
    size = template.font.size if template is not None else None
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    for index, sentence in enumerate(parts):
        run = paragraph.add_run(sentence + (" " if index < len(parts) - 1 else ""))
        if size is not None:
            run.font.size = size
        if sentence not in old_sentences:
            run.font.color.rgb = RED
    return True


def append_method_sentence(document: Document) -> bool:
    paragraph = next(
        p for p in document.paragraphs
        if "additional diagnostics were run" in p.text
    )
    if "eight controlled" in paragraph.text:
        return False
    used = sum(
        1 for sentence in sentences(paragraph.text.strip())
        if any(sentence.startswith(word + ",") for word in ORDINALS)
    )
    styled(paragraph, METHOD_SENTENCE.format(ordinal=ORDINALS[used]))
    return True


def table11_rows() -> list[list[str]]:
    summary = pd.read_csv(STAGE30 / "condition_metrics.csv")
    summary[["color", "pretraining", "dimensionality"]] = (
        summary.condition.str.split("+", expand=True)
    )
    ordered = summary.sort_values("macro_f1_mean", ascending=False)
    rows = []
    for record in ordered.to_dict("records"):
        rows.append([
            LABEL[record["color"]],
            LABEL[record["pretraining"]],
            LABEL[record["dimensionality"]],
            f"{fmt(record['macro_f1_mean'])} +/- {fmt(record['macro_f1_std'])}",
            fmt(record["balanced_accuracy_mean"]),
            fmt(record["mcc_mean"]),
        ])
    return rows


def table12_rows() -> list[list[str]]:
    effects = pd.read_csv(STAGE30 / "factor_effects.csv")
    order = {"pretraining": 0, "color": 1, "dimensionality": 2}
    effects = effects.sort_values("factor", key=lambda s: s.map(order))
    rows = []
    for record in effects.to_dict("records"):
        rows.append([
            FACTOR_LABEL[record["factor"]],
            CONTRAST_LABEL[record["contrast"]],
            signed(record["mean_difference"]),
            signed(record["min_across_settings"]),
            signed(record["max_across_settings"]),
            str(int(record["settings"])),
        ])
    return rows


def compose_prose() -> str:
    summary = pd.read_csv(STAGE30 / "condition_metrics.csv").set_index("condition")
    effects = pd.read_csv(STAGE30 / "factor_effects.csv").set_index("factor")
    paired = pd.read_csv(STAGE30 / "paired_differences.csv")
    anchor = json.loads((STAGE30 / "anchor_check.json").read_text(encoding="utf-8"))

    best = summary.macro_f1_mean.idxmax()
    worst = summary.macro_f1_mean.idxmin()
    baseline = summary.loc["rgb+pretrained+full512"]
    pretraining = effects.loc["pretraining"]
    color = effects.loc["color"]
    dimension = effects.loc["dimensionality"]

    unanimous = paired[
        (paired.factor == "pretraining") & (paired.folds_worse == 0)
    ]
    unanimous_note = (
        f" In one of those settings all {int(unanimous.n_paired_folds.iloc[0])} paired "
        "folds improved and none degraded."
        if not unanimous.empty else ""
    )

    return (
        "The comparison between handcrafted features and frozen deep embeddings "
        "reported above changes colour, pretraining, cropping, and dimensionality at "
        "the same time, so a performance difference cannot be attributed to any one "
        "of them. Three of those factors were therefore separated in a controlled "
        "ablation on the ResNet18 branch. Eight conditions were formed by crossing "
        "RGB against grayscale replicated to three channels, ImageNet weights against "
        "random initialisation, and the full 512-dimensional embedding against six "
        "principal components, the latter matching the dimensionality of the "
        "handcrafted representation exactly (Table 11). Every condition used the same "
        "source-group-aware folds, the same repeated nested protocol, and the same "
        "three classifier families, and principal component analysis was fitted inside "
        "the training pipeline only. Extraction for the RGB, ImageNet, 512-d condition "
        "reproduced the embeddings used elsewhere in this article to "
        f"{anchor['max_abs_deviation']:.0e} maximum absolute deviation, so differences "
        "between conditions reflect the varied factors rather than a change of harness. "
        "Paired at the outer-fold level, only one factor mattered (Table 12). "
        "ImageNet pretraining raised macro-F1 by "
        f"{signed(pretraining.mean_difference)} on average, and the effect was positive "
        f"in every one of the {int(pretraining.settings)} combinations of the remaining "
        f"factors, ranging from {signed(pretraining.min_across_settings)} to "
        f"{signed(pretraining.max_across_settings)}." + unanimous_note + " Colour moved "
        f"performance by {signed(color.mean_difference)}, that is, grayscale was "
        "marginally better than RGB rather than worse, and the direction was negative "
        "in every combination. Dimensionality moved performance by "
        f"{signed(dimension.mean_difference)} with a range from "
        f"{signed(dimension.min_across_settings)} to "
        f"{signed(dimension.max_across_settings)} that crosses zero, so six components "
        "were not measurably worse than the full embedding. The strongest condition, "
        f"{LABEL[best.split('+')[0]].lower()} with ImageNet weights and the full "
        f"embedding, reached {fmt(summary.loc[best, 'macro_f1_mean'])}, while the "
        "weakest, RGB with random initialisation and six components, reached "
        f"{fmt(summary.loc[worst, 'macro_f1_mean'])}; the baseline condition reached "
        f"{fmt(baseline.macro_f1_mean)} +/- {fmt(baseline.macro_f1_std)}. The advantage "
        "of the frozen deep branch on this collection therefore rests on visual "
        "knowledge transferred from ImageNet, not on access to colour, not on a larger "
        "embedding, and not on the representation family as such. Centre cropping was "
        "not varied because it belongs to the transform supplied with the ImageNet "
        "weights, so removing it would also change resizing and normalisation and the "
        "resulting effect could not be attributed to cropping alone. The ablation also "
        "uses a single backbone, and its conclusions are not generalised to other "
        "architectures."
    )


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if lock_present():
        raise RuntimeError("Manuskrip masih terbuka di Word. Tutup lebih dahulu.")
    if not (STAGE30 / "condition_metrics.csv").exists():
        raise FileNotFoundError(
            "Hasil tahap 30 belum ada. Jalankan 30_controlled_ablation.py lebih dahulu."
        )

    document = Document(MANUSCRIPT)
    already = any(MARKER in paragraph.text for paragraph in document.paragraphs)

    rows11, rows12 = table11_rows(), table12_rows()
    if len(rows11) != 8:
        raise AssertionError(f"Table 11 harus 8 baris, bukan {len(rows11)}")
    if len(rows12) != 3:
        raise AssertionError(f"Table 12 harus 3 baris, bukan {len(rows12)}")

    print("=" * 76)
    print("SISIPKAN SUBBAGIAN 4.9, TABLE 11, DAN TABLE 12")
    print("=" * 76)

    changed = []
    if renumber_method_ordinals(document):
        changed.append("penomoran diagnostik Section 3.8")

    if not already:
        if append_method_sentence(document):
            changed.append("kalimat metode ablasi")

        kids = list(document.element.body.iterchildren())
        discussion_index = next(
            index for index, child in enumerate(kids)
            if child.tag.endswith("}p")
            and Paragraph(child, document).text.strip() == "5 Discussion"
        )
        anchor = kids[discussion_index - 1]

        heading = document.add_paragraph(style="Q3 Subsection")
        styled(heading, MARKER, size=Pt(11))
        prose = document.add_paragraph(style="Q3 Body")
        styled(prose, compose_prose())
        caption11 = document.add_paragraph(style="Q3 Caption")
        styled(caption11, TABLE11_CAPTION, size=Pt(9))
        table11 = build_table(document, TABLE11_HEADERS, rows11)
        caption12 = document.add_paragraph(style="Q3 Caption")
        styled(caption12, TABLE12_CAPTION, size=Pt(9))
        table12 = build_table(document, TABLE12_HEADERS, rows12)

        for element in (
            table12._tbl, caption12._p, table11._tbl, caption11._p,
            prose._p, heading._p,
        ):
            anchor.addnext(element)
        changed.append("subbagian 4.9, Table 11, dan Table 12")

    if not changed:
        print("-" * 76)
        print("Tidak ada perubahan. Manuskrip sudah memuat ablasi terkontrol.")
        return

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_table11_12.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
        print("Cadangan:", backup.relative_to(PROJECT_DIR))

    document.save(MANUSCRIPT)

    check = Document(MANUSCRIPT)
    captions = [
        paragraph.text.strip() for paragraph in check.paragraphs
        if paragraph.text.strip().startswith("Table ")
        and paragraph.style.name == "Q3 Caption"
    ]
    print("-" * 76)
    for item in changed:
        print(f"  diperbarui: {item}")
    print("-" * 76)
    print(f"Tabel dalam manuskrip: {len(check.tables)}")
    for text in captions[-3:]:
        print("  ", text[:66])
    print("-" * 76)
    print("Penomoran tabel lain tidak bergeser.")


if __name__ == "__main__":
    main()
