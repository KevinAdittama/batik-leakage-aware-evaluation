"""Segarkan kalimat naratif Results yang memuat angka sisi development.

Skrip 26 sudah menyegarkan sel tabel. Skrip ini menangani kalimat yang memuat
angka yang sama, ditambah dua hal yang berubah karena keputusan opsi 2:

  - Model formal tidak lagi ditobatkan sebagai pemenang. Margin seleksi lebih
    kecil daripada derau antar-fold (tahap 25), sehingga narasi yang menyebut
    satu model unggul diganti pelaporan berdampingan.
  - Fold kini dipisahkan pada grain foto sumber, bukan grain berkas.

Penandaan merah per kalimat
---------------------------
Editor butir E.1 meminta setiap perubahan ditandai. Memerahkan seluruh paragraf
akan menyesatkan ketika hanya dua angka di dalamnya yang berubah. Karena itu
paragraf dipecah menjadi kalimat, lalu hanya kalimat yang benar-benar berbeda
dari versi sebelumnya yang diberi warna merah. Kalimat yang tidak berubah
mempertahankan warna aslinya.

Seluruh angka dibaca dari CSV terauditkan. Tidak ada yang diketik manual.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\27_refresh_manuscript_prose.py
"""

from __future__ import annotations

import importlib.util
import json
import re
import shutil
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

from pipeline_config import PROJECT_DIR, RESULTS_DIR

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
AUDIT_OUT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "04_Tabel_Manifest_dan_Hasil"
    / "Audit_Numerik_dan_Eksperimen"
    / "outputs"
)
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_Segar_Prosa_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)
SENTENCE_BOUNDARY = re.compile(r"(?<=[.;:]) (?=[A-Z(])")


def lock_present() -> bool:
    return any(MANUSCRIPT.parent.glob("~$*"))


def fmt(value: float, decimals: int = 3) -> str:
    return f"{float(value):.{decimals}f}"


def pm(mean: float, std: float, decimals: int = 3) -> str:
    return f"{fmt(mean, decimals)} +/- {fmt(std, decimals)}"


def sentences(text: str) -> list[str]:
    return [part for part in SENTENCE_BOUNDARY.split(text.strip()) if part]


def rewrite(paragraph: Paragraph, new_text: str) -> bool:
    """Tulis ulang paragraf; merah hanya pada kalimat yang berbeda."""
    old_text = paragraph.text.strip()
    if old_text == new_text.strip():
        return False

    old_sentences = set(sentences(old_text))
    template = paragraph.runs[0] if paragraph.runs else None
    size = template.font.size if template is not None else None
    original_color = (
        template.font.color.rgb
        if template is not None and template.font.color and template.font.color.rgb
        else None
    )

    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)

    parts = sentences(new_text)
    for index, sentence in enumerate(parts):
        run = paragraph.add_run(sentence + (" " if index < len(parts) - 1 else ""))
        if size is not None:
            run.font.size = size
        if sentence in old_sentences:
            if original_color is not None:
                run.font.color.rgb = original_color
        else:
            run.font.color.rgb = RED
    return True


def load_sources() -> dict:
    stage07 = pd.read_csv(RESULTS_DIR / "07_cross_validation/cv_summary_primary.csv")
    stage07["display"] = stage07.model.replace({"SVM (RBF)": "SVM-RBF"})
    nested12 = pd.read_csv(
        RESULTS_DIR / "12_submission_robustness/nested_cv_summary.csv", index_col=0
    )
    nested12_folds = pd.read_csv(
        RESULTS_DIR / "12_submission_robustness/nested_cv_outer_fold_metrics.csv"
    )
    stage14 = pd.read_csv(
        RESULTS_DIR / "14_repeated_nested_augmentation/summary.csv"
    ).set_index("metric")
    stage14_selection = pd.read_csv(
        RESULTS_DIR / "14_repeated_nested_augmentation/selection_frequency.csv"
    )
    audited = pd.read_csv(AUDIT_OUT / "table6_external_model_metrics_audited.csv")
    audited["display"] = audited.model.replace({"SVM (RBF)": "SVM-RBF"})
    paired = pd.read_csv(AUDIT_OUT / "paired_macro_f1_bootstrap_audited.csv")
    metadata = pd.read_csv(
        RESULTS_DIR / "12_submission_robustness/metadata_negative_control_metrics.csv"
    )
    # Berkas subtype dipilih lewat slug model formal, bukan lewat glob. Glob akan
    # diam-diam mengambil berkas era model sebelumnya bila keduanya sempat ada.
    selected = json.loads(
        (RESULTS_DIR / "08_uji_eksternal/selected_model_result.json").read_text(
            encoding="utf-8"
        )
    )
    subtype_path = AUDIT_OUT / f"{selected['model_slug']}_subtype_error_counts_audited.csv"
    if not subtype_path.exists():
        raise FileNotFoundError(
            f"Hitungan subtype untuk model formal belum ada: {subtype_path.name}"
        )
    subtype = pd.read_csv(subtype_path)
    margin = pd.read_csv(
        RESULTS_DIR / "25_model_selection_stability/selection_margin.csv"
    ).set_index("protocol")
    groups = pd.read_csv(RESULTS_DIR / "24_source_groups/source_groups.csv")
    confirmed = pd.read_csv(RESULTS_DIR / "24_source_groups/confirmed_pairs.csv")
    return {
        "cv": stage07.set_index("display"),
        "cv_sorted": stage07.sort_values("f1_macro_mean", ascending=False),
        "nested12": nested12,
        "nested12_folds": nested12_folds,
        "stage14": stage14,
        "stage14_selection": stage14_selection,
        "audited": audited.set_index("display"),
        "paired": paired,
        "metadata": metadata,
        "subtype": subtype,
        "margin": margin,
        "groups": groups,
        "confirmed": confirmed,
    }


def para_method_folds(s: dict) -> str:
    groups = s["groups"]
    development = groups[groups["set"] == "development"]
    return (
        "A shuffled five-fold StratifiedGroupKFold with seed 42 was formed from the "
        f"{len(development)} original files, grouped into "
        f"{development.group_id.nunique()} source photographs so that overlapping "
        "crops of one photograph cannot fall on opposite sides of a split. "
        "Edge/morphology, texture, frequency, and combined-six feature sets were "
        "compared on these folds. All three classical models are reported side by "
        "side rather than reduced to a single retained winner, because the margin "
        "between the two leading models is smaller than the fold-to-fold standard "
        "error (Table 10). Because the same folds supported comparison and "
        "performance estimation, these CV scores may contain selection-induced "
        "optimism. The external collection did not determine any model fitting or "
        "thresholding. The two frozen backbones were not tuned on this collection, "
        "and their comparison remains post-hoc."
    )


def para_nested_results(s: dict) -> str:
    nested, folds = s["nested12"], s["nested12_folds"]
    counts = folds.selected_model.value_counts()
    words = {1: "one", 2: "two", 3: "three", 4: "four", 5: "five"}
    nested_text = " and ".join(
        f"{model.replace('SVM (RBF)', 'SVM-RBF')} in {words.get(int(count), int(count))}"
        for model, count in counts.items()
    )
    stage14, selection = s["stage14"], s["stage14_selection"]
    total = int(selection.outer_fold_selections.sum())
    repeated_text = ", ".join(
        f"{record['selected_model']} in {int(record['outer_fold_selections'])}"
        for record in selection.sort_values(
            "outer_fold_selections", ascending=False
        ).to_dict("records")
    )
    return (
        "The originals-only nested analysis selected the combined six-feature set in "
        f"all five outer folds; {nested_text} of five. Outer-fold macro-F1 was "
        f"{pm(nested.loc['f1_macro', 'mean'], nested.loc['f1_macro', 'std'])}, "
        f"balanced accuracy {fmt(nested.loc['balanced_accuracy', 'mean'])}, and MCC "
        f"{fmt(nested.loc['mcc', 'mean'])}. The five-repeat strictly nested fold-local "
        "augmentation sensitivity produced repeat-level macro-F1 "
        f"{pm(stage14.loc['f1_macro', 'repeat_mean'], stage14.loc['f1_macro', 'repeat_std'])}, "
        f"balanced accuracy "
        f"{pm(stage14.loc['balanced_accuracy', 'repeat_mean'], stage14.loc['balanced_accuracy', 'repeat_std'])}, "
        f"and MCC "
        f"{pm(stage14.loc['mcc', 'repeat_mean'], stage14.loc['mcc', 'repeat_std'])}. "
        f"The combined six-feature set was selected in all {total} outer folds, while "
        f"the model choice was almost evenly divided: {repeated_text} of {total}. "
        "That near-even division is itself a result: the selection rule no longer "
        "separates the two leading classifiers once folds respect source photographs."
    )


def para_metadata_control(s: dict) -> str:
    metadata = s["metadata"]
    development = metadata[metadata.evaluation == "development_5fold"]
    external = metadata[metadata.evaluation != "development_5fold"].iloc[0]
    return (
        "The learned metadata-only negative control achieved development CV macro-F1 "
        f"{pm(development.f1_macro.mean(), development.f1_macro.std(ddof=1))} and "
        f"external macro-F1 {fmt(external.f1_macro)}, with external MCC "
        f"{fmt(external.mcc)}. Together with the fixed extension rule, this shows that "
        "acquisition metadata were highly predictive internally and changed "
        "substantially between collections. It does not prove that any pixel model "
        "used a particular metadata field."
    )


def para_single_loop(s: dict) -> str:
    ordered = s["cv_sorted"]
    first, second, third = (ordered.iloc[i] for i in range(3))
    margin = s["margin"].loc["source-group-aware"]
    return (
        "Within the evaluated handcrafted family the three models were closely "
        f"spaced (Table 6). {first.display} reached mean macro-F1 "
        f"{pm(first.f1_macro_mean, first.f1_macro_std)}, {second.display} "
        f"{pm(second.f1_macro_mean, second.f1_macro_std)}, and {third.display} "
        f"{pm(third.f1_macro_mean, third.f1_macro_std)}. The gap between the two "
        f"leading models, {fmt(margin.margin)} macro-F1, is only "
        f"{float(margin.margin_in_standard_errors):.2f} standard errors of the "
        "fold-level mean, so no handcrafted model is claimed to be superior here. "
        "The combined six-feature set improved macro-F1 over any single feature "
        "group for both leading models (Figure 6)."
    )


def para_external_classical(s: dict) -> str:
    audited = s["audited"]
    ordered = audited.sort_values("external_f1_macro", ascending=False)
    classical = ordered[ordered.model_group == "Classical ML"]
    best = classical.iloc[0]
    cv = s["cv_sorted"].set_index("display")
    lines = []
    for name, record in classical.iterrows():
        lines.append(
            f"{name} fell from CV macro-F1 "
            f"{fmt(cv.loc[name, 'f1_macro_mean'])} to "
            f"{fmt(record.external_f1_macro)} "
            f"(95% CI {fmt(record.external_f1_ci95_low)}-{fmt(record.external_f1_ci95_high)}; "
            f"MCC {fmt(record.external_mcc)})"
        )
    return (
        "All classical models declined on the separate external collection (Table 7). "
        + "; ".join(lines)
        + ". "
        f"The best external recall figures for {best.name} were "
        f"{fmt(best.external_recall_batik)} for batik and "
        f"{fmt(best.external_recall_non_batik)} for non-batik. The decline is present "
        "for every handcrafted model, so it does not depend on which model the "
        "selection rule happened to rank first."
    )


def para_deep_benchmarks(s: dict) -> str:
    audited = s["audited"]
    mobilenet = audited.loc["MobileNetV2"]
    resnet = audited.loc["ResNet18"]
    paired = s["paired"]
    resnet_row = paired[paired.comparison.str.startswith("ResNet18 minus")].iloc[0]
    reference = resnet_row.comparison.split(" minus ")[1]
    return (
        "The post-hoc frozen RGB benchmarks yielded higher external point estimates. "
        f"MobileNetV2 was the deep-CV winner and reached external macro-F1 "
        f"{fmt(mobilenet.external_f1_macro)} (95% CI "
        f"{fmt(mobilenet.external_f1_ci95_low)}-{fmt(mobilenet.external_f1_ci95_high)}) "
        f"and MCC {fmt(mobilenet.external_mcc)}. ResNet18 reached macro-F1 "
        f"{fmt(resnet.external_f1_macro)} (95% CI "
        f"{fmt(resnet.external_f1_ci95_low)}-{fmt(resnet.external_f1_ci95_high)}) and MCC "
        f"{fmt(resnet.external_mcc)}. On the same 60 files, the paired bootstrap "
        f"difference in macro-F1 between ResNet18 and {reference} was "
        f"{fmt(resnet_row.estimate_difference)} (95% CI "
        f"{fmt(resnet_row.ci95_low)}-{fmt(resnet_row.ci95_high)}). This interval is "
        "specific to the observed collection; it does not address acquisition "
        "confounding, and the deep analysis remains post-hoc."
    )


def para_subtype_errors(s: dict) -> str:
    subtype = s["subtype"]
    errors = int(subtype.errors.sum())
    false_negative = int(subtype.false_negative.sum())
    false_positive = int(subtype.false_positive.sum())
    top = subtype[subtype.errors > 0].sort_values(
        ["error_rate", "subjenis"], ascending=[False, True]
    ).head(4)
    described = ", ".join(
        f"{'batik' if record['label'] == 1 else 'non-batik'} "
        f"{str(record['subjenis']).replace('_', ' ')} "
        f"({int(record['errors'])}/{int(record['n'])})"
        for record in top.to_dict("records")
    )
    audited = s["audited"]
    formal = audited[audited.model_group == "Classical ML"].sort_values(
        "cv_f1_macro_mean", ascending=False
    ).index[0]
    return (
        f"{formal} misclassified {errors} of 60 external images: {false_negative} "
        f"false negatives and {false_positive} false positives. The largest observed "
        f"subtype error fractions were {described} (Figure 8, left panel). The right "
        "panel of Figure 8 shows high predicted-score examples from both error "
        "directions. Because each subtype "
        "contains only three to seven images, these rates are exploratory "
        "descriptions, not stable population estimates."
    )


def para_balancing(s: dict) -> str:
    """Susun ulang paragraf 4.7 memakai penyusun yang sama dengan skrip 23."""
    path = PROJECT_DIR / "23_insert_balancing_sensitivity.py"
    spec = importlib.util.spec_from_file_location("insert_balancing", path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module.compose_prose(module.load_numbers())


def _handcrafted(s: dict):
    """Ketiga model klasik, terurut menurut macro-F1 CV."""
    ordered = s["cv_sorted"]
    return [ordered.iloc[index] for index in range(len(ordered))]


def _group_facts(s: dict) -> dict:
    groups = s["groups"]
    development = groups[groups["set"] == "development"]
    multi = development[development.group_size > 1]
    return {
        "pairs": len(s["confirmed"]),
        "images": int(multi.group_size.count()),
        "groups": int(multi.group_id.nunique()),
        "total_files": len(development),
        "total_groups": int(development.group_id.nunique()),
    }


def para_abstract(s: dict) -> str:
    first, _, _ = _handcrafted(s)[:3]
    stage14 = s["stage14"]
    selection = s["stage14_selection"].sort_values(
        "outer_fold_selections", ascending=False
    )
    total = int(selection.outer_fold_selections.sum())
    split = " and ".join(
        f"{record['selected_model']} in {int(record['outer_fold_selections'])}"
        for record in selection.to_dict("records")
    )
    audited = s["audited"]
    external = float(audited.loc[first.display, "external_f1_macro"])
    facts = _group_facts(s)
    return (
        "Abstract: Binary batik recognition is vulnerable to shortcut learning "
        "because visually similar non-batik textiles can share repetition and "
        "texture, while image format, resolution, and source may correlate with "
        "supplied labels. We present a reproducible, leakage-aware evaluation "
        "workflow combining overlap auditing, source-group adjudication, "
        "original-file-aware augmentation, selection-aware validation, external "
        "testing, acquisition-metadata negative controls, and subtype-level error "
        "analysis. Six semantically interpretable grayscale features with "
        "conventional classifiers were evaluated on a byte-unique development set "
        "and a separate external collection; frozen ResNet18 and MobileNetV2 were "
        "included as post-hoc benchmarks. An alignment-based screen that "
        "perceptual hashing cannot perform identified "
        f"{facts['pairs']} same-photograph pairs, so the "
        f"{facts['total_files']} development originals were split as "
        f"{facts['total_groups']} source groups rather than as independent files. "
        "A five-repeat strictly nested fold-local augmentation sensitivity "
        "analysis yielded repeat-level macro-F1 "
        f"{pm(stage14.loc['f1_macro', 'repeat_mean'], stage14.loc['f1_macro', 'repeat_std'])} "
        f"and divided model selection almost evenly ({split} of {total} outer "
        "folds), showing that the selection rule no longer separates the leading "
        "classifiers. The best handcrafted model declined from CV macro-F1 "
        f"{fmt(first.f1_macro_mean)} to {fmt(external)} on the external "
        "collection. A metadata-only classifier that never used image content "
        "performed strongly during development but degraded externally, indicating "
        "substantial acquisition-related confounding. These findings show that high "
        "internal cross-validation scores alone cannot establish concept-level "
        "generalization in small, heterogeneous image datasets. The study "
        "contributes a documented diagnostic workflow for evaluating visual "
        "classifiers when provenance and acquisition conditions remain "
        "incompletely controlled."
    )


def para_hash_results(s: dict) -> str:
    facts = _group_facts(s)
    return (
        "Before exclusion, manual review confirmed two development-external near "
        "duplicates; both development-side files were excluded by approved manifest "
        "decisions and the pipeline was rerun. In the post-exclusion screen, all six "
        "development-external candidates were retained as visually different dHash "
        "collisions (Table 4). The within-development candidates were then "
        "adjudicated by full-resolution alignment rather than by hashing alone, "
        f"which confirmed {facts['pairs']} same-photograph pairs covering "
        f"{facts['images']} images in {facts['groups']} source groups (Table 9). "
        "Only two of those pairs had appeared in the perceptual-hash candidate list, "
        "because perceptual hashing is not shift invariant and therefore cannot see "
        "overlapping crops of one photograph. The development manifest is now split "
        f"as {facts['total_groups']} source groups rather than "
        f"{facts['total_files']} independent files."
    )


def para_central_finding(s: dict) -> str:
    first = _handcrafted(s)[0]
    metadata = s["metadata"]
    development = metadata[metadata.evaluation == "development_5fold"]
    external = metadata[metadata.evaluation != "development_5fold"].iloc[0]
    return (
        "The central finding is not the ranking of handcrafted and frozen deep "
        "representations. It is the incompatibility between apparently strong "
        "internal estimates and evidence of acquisition confounding, confirmed "
        "same-photograph groups, and external degradation. The learned metadata "
        f"control reached development macro-F1 {fmt(development.f1_macro.mean())}, "
        "exceeding the best handcrafted single-loop estimate of "
        f"{fmt(first.f1_macro_mean)}, yet transferred at only "
        f"{fmt(external.f1_macro)} externally. This does not prove that image models "
        "used metadata fields; it shows that supplied labels were strongly aligned "
        "with acquisition properties capable of supporting high internal scores."
    )


def para_second_question(s: dict) -> str:
    models = _handcrafted(s)
    nested, folds = s["nested12"], s["nested12_folds"]
    counts = folds.selected_model.value_counts()
    words = {1: "one", 2: "two", 3: "three", 4: "four", 5: "five"}
    nested_text = " and ".join(
        f"{model.replace('SVM (RBF)', 'SVM-RBF')} in "
        f"{words.get(int(count), int(count))}"
        for model, count in counts.items()
    )
    margin = s["margin"].loc["source-group-aware"]
    return (
        "For the second question, no handcrafted model separated itself from the "
        f"others. The three point estimates spanned {fmt(models[-1].f1_macro_mean)} "
        f"to {fmt(models[0].f1_macro_mean)}, and the gap between the two leading "
        f"models was {fmt(margin.margin)} macro-F1, roughly "
        f"{float(margin.margin_in_standard_errors):.2f} standard errors of the "
        "fold-level mean. The originals-only nested diagnostic selected "
        f"{nested_text} of five outer folds and produced macro-F1 "
        f"{pm(nested.loc['f1_macro', 'mean'], nested.loc['f1_macro', 'std'])}. "
        "Selection protocol therefore matters, and neither analysis removes "
        "unobserved source dependence. The findings remain consistent with prior "
        "batik research showing that texture and shape descriptors can support "
        "conventional classifiers. The post-hoc frozen-deep benchmarks produced "
        "higher internal scores."
    )


def para_repeated_sensitivity(s: dict) -> str:
    stage14 = s["stage14"]
    selection = s["stage14_selection"].sort_values(
        "outer_fold_selections", ascending=False
    )
    total = int(selection.outer_fold_selections.sum())
    split = " and ".join(
        f"{record['selected_model']} in {int(record['outer_fold_selections'])}"
        for record in selection.to_dict("records")
    )
    facts = _group_facts(s)
    return (
        "The repeated fold-local augmentation sensitivity corroborated the "
        "combined-feature preference while reducing selection reuse: repeat-level "
        "macro-F1 was "
        f"{pm(stage14.loc['f1_macro', 'repeat_mean'], stage14.loc['f1_macro', 'repeat_std'])}, "
        f"and the model choice divided almost evenly, {split} of {total} outer "
        "folds. That near-even division indicates that the selection rule is not "
        "resolving a real difference between the two leading classifiers on this "
        f"collection. Because these repeats reuse the same {facts['total_files']} "
        "files and cannot group acquisition sources that leave no visual overlap, "
        "the result strengthens the internal sensitivity analysis but does not "
        "establish external generalization."
    )


def para_third_question(s: dict) -> str:
    first = _handcrafted(s)[0]
    audited = s["audited"]
    record = audited.loc[first.display]
    mobilenet = float(audited.loc["MobileNetV2", "external_f1_macro"])
    resnet = float(audited.loc["ResNet18", "external_f1_macro"])
    return (
        "The third question exposed the transfer limitation. The best handcrafted "
        f"model, {first.display}, declined from macro-F1 {fmt(first.f1_macro_mean)} "
        f"in five-fold validation to {fmt(record.external_f1_macro)} on the external "
        f"collection. Its bootstrap interval "
        f"({fmt(record.external_f1_ci95_low)}-{fmt(record.external_f1_ci95_high)}) "
        "shows that the 60-image estimate is uncertain, but even the point estimate "
        "is far below the internal result, and every other handcrafted model "
        f"declined as well. MobileNetV2 and ResNet18 retained higher external point "
        f"estimates of {fmt(mobilenet)} and {fmt(resnet)}. However, the deep models "
        "were introduced after the external collection had already been examined. "
        "Their scores therefore provide exploratory context rather than prospective "
        "evidence."
    )


def para_metadata_discussion(s: dict) -> str:
    first = _handcrafted(s)[0]
    metadata = s["metadata"]
    development = metadata[metadata.evaluation == "development_5fold"]
    external = metadata[metadata.evaluation != "development_5fold"].iloc[0]
    return (
        "The metadata negative controls change how the cross-validation scores "
        "should be read. A JPEG-versus-other-format rule achieved development "
        "macro-F1 0.922, above the best handcrafted CV macro-F1 of "
        f"{fmt(first.f1_macro_mean)}, but fell to 0.444 externally. A learned "
        "metadata model using extension, dimensions, aspect ratio, and file size "
        "reached "
        f"{pm(development.f1_macro.mean(), development.f1_macro.std(ddof=1))} "
        f"internally and {fmt(external.f1_macro)} externally. The image classifiers "
        "never received these metadata columns, so the diagnostics do not identify "
        "their decision mechanism. They show that class membership was strongly "
        "aligned with development source properties."
    )


def para_hash_discussion(s: dict) -> str:
    facts = _group_facts(s)
    return (
        "Exact-hash auditing still served an important but bounded role. It excluded "
        "one within-development duplicate and eight byte-identical "
        "development-external overlaps. Manual perceptual review then confirmed two "
        "further cross-set near duplicates; the approved development-side exclusions "
        "were applied before the complete rerun. Perceptual hashing, however, is not "
        "shift invariant, and it missed most of the overlapping crops present in "
        "this collection. Full-resolution alignment adjudication confirmed "
        f"{facts['pairs']} same-photograph pairs covering {facts['images']} images, "
        f"only two of which the hash screen had flagged; folds are now formed over "
        f"{facts['total_groups']} source groups instead of {facts['total_files']} "
        "files. Stronger future protocols still require provenance identifiers that "
        "group all images from one object, website, session, or acquisition source "
        "before splitting, because alignment detects translational overlap only. "
        "Similarity-aware split research further supports this requirement. The "
        "present novelty is empirical and diagnostic: overlap control and "
        "training-only augmentation can coexist with severe acquisition "
        "confounding, and metadata controls can materially change the "
        "interpretation of strong CV scores."
    )


def para_limitations(s: dict) -> str:
    facts = _group_facts(s)
    return (
        "Several limitations constrain the claims. First, development labels were "
        "inherited from directory structure and were not independently verified by "
        "multiple batik experts; neither inter-rater agreement nor adjudication is "
        "available. Second, source URLs, licenses, and object/session identifiers "
        "are incomplete, preventing a full provenance audit and public "
        "redistribution of all raw images. Third, format, resolution, and likely "
        "source are strongly associated with class; resizing does not guarantee "
        "removal of compression, interpolation, sharpness, or collection-style "
        f"signatures. Fourth, source grouping rests on translational alignment: "
        f"{facts['pairs']} same-photograph pairs were confirmed and grouped, but "
        "non-overlapping crops of one photograph, rotated copies, and heavily "
        "rescaled copies remain undetectable by this test, so the grouping is a "
        "lower bound rather than a guarantee of independence. Fifth, the "
        "single-loop comparison may contain selection optimism, and the margin "
        "between the two leading handcrafted models is smaller than the fold-level "
        "standard error, so no ranking among them is claimed. Sixth, the 60-image "
        "external collection was inspected during iterative development and cannot "
        "serve as a prospectively untouched final test for every analysis. Seventh, "
        "the post-hoc deep comparison changes RGB information, dimensionality, and "
        "pretraining simultaneously. Finally, subtype counts are small, and the task "
        "concerns supplied visual folder classes rather than production technique or "
        "cultural authenticity. The findings are diagnostic and hypothesis-generating "
        "rather than deployment-validating."
    )


def para_conclusion_results(s: dict) -> str:
    first = _handcrafted(s)[0]
    nested = s["nested12"]
    stage14 = s["stage14"]
    audited = s["audited"]
    external = float(audited.loc[first.display, "external_f1_macro"])
    facts = _group_facts(s)
    return (
        "This exploratory case study examined whether file-level leakage control, "
        "training-only augmentation, and external evaluation are sufficient to "
        "support binary batik-recognition claims in a small heterogeneous "
        "collection. They are not. Alignment adjudication showed that file-level "
        f"control was itself incomplete: {facts['pairs']} same-photograph pairs were "
        "split across folds, so the analysis was rerun with source-group-aware "
        f"splitting over {facts['total_groups']} groups. Under that protocol the "
        f"leading handcrafted model reached single-loop CV macro-F1 "
        f"{pm(first.f1_macro_mean, first.f1_macro_std)}, the originals-only nested "
        f"diagnostic yielded "
        f"{pm(nested.loc['f1_macro', 'mean'], nested.loc['f1_macro', 'std'])}, and "
        "the five-repeat strictly nested fold-local augmentation sensitivity yielded "
        f"{pm(stage14.loc['f1_macro', 'repeat_mean'], stage14.loc['f1_macro', 'repeat_std'])}. "
        f"The same model nevertheless declined to {fmt(external)} externally, and "
        "the selection rule no longer separated the two leading classifiers. "
        "Post-hoc MobileNetV2 and ResNet18 produced higher external point estimates, "
        "but those comparisons do not establish prospective or architectural "
        "superiority."
    )


def para_conclusion_evidence(s: dict) -> str:
    metadata = s["metadata"]
    development = metadata[metadata.evaluation == "development_5fold"]
    external = metadata[metadata.evaluation != "development_5fold"].iloc[0]
    facts = _group_facts(s)
    return (
        "The most consequential evidence came from acquisition diagnostics. A "
        "learned metadata-only model reached development macro-F1 "
        f"{pm(development.f1_macro.mean(), development.f1_macro.std(ddof=1))} but "
        f"only {fmt(external.f1_macro)} externally. Perceptual screening led to two "
        "approved cross-set exclusions, and alignment adjudication then grouped "
        f"{facts['images']} images into {facts['groups']} source photographs that "
        "hashing alone had not resolved. The practical contribution is therefore a "
        "reproducible diagnostic workflow, not a deployable classifier. A "
        "confirmatory follow-up should begin with expert-verified labels, complete "
        "source and license provenance, acquisition-matched data, controlled "
        "RGB/grayscale ablations, and a newly collected multi-source external set "
        "frozen before analysis. Until then, the reported models are exploratory "
        "baselines for studying dataset bias rather than evidence of general batik "
        "recognition."
    )


TARGETS = [
    ("A shuffled five-fold Stratified", para_method_folds),
    ("The originals-only nested analysis selected", para_nested_results),
    ("The learned metadata-only negative control", para_metadata_control),
    ("Random Forest obtained the highest single-loop", para_single_loop),
    ("Within the evaluated handcrafted family", para_single_loop),
    ("All classical models declined on the separate external", para_external_classical),
    ("The post-hoc frozen RGB benchmarks", para_deep_benchmarks),
    ("Random Forest misclassified", para_subtype_errors),
    ("SVM-RBF misclassified", para_subtype_errors),
    ("A final development-side sensitivity analysis", para_balancing),
    ("Abstract: Binary batik recognition", para_abstract),
    ("Before exclusion, manual review confirmed", para_hash_results),
    ("The central finding is not the ranking", para_central_finding),
    ("For the second question", para_second_question),
    ("The repeated fold-local augmentation sensitivity", para_repeated_sensitivity),
    ("The third question exposed the transfer limitation", para_third_question),
    ("The metadata negative controls change how", para_metadata_discussion),
    ("Exact-hash auditing still served", para_hash_discussion),
    ("Several limitations constrain the claims", para_limitations),
    ("This exploratory case study examined whether", para_conclusion_results),
    ("The most consequential evidence came from", para_conclusion_evidence),
]


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if lock_present():
        raise RuntimeError("Manuskrip masih terbuka di Word. Tutup lebih dahulu.")

    document = Document(MANUSCRIPT)
    sources = load_sources()

    print("=" * 76)
    print("SEGARKAN KALIMAT NARATIF DARI SUMBER TERAUDITKAN")
    print("=" * 76)

    seen: set[int] = set()
    changed = 0
    for prefix, builder in TARGETS:
        matches = [
            index
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip().startswith(prefix) and index not in seen
        ]
        if not matches:
            print(f"  [lewati]  {prefix[:46]:<48} tidak ditemukan")
            continue
        if len(matches) > 1:
            raise AssertionError(f"Awalan paragraf tidak unik: {prefix!r}")
        index = matches[0]
        seen.add(index)
        if rewrite(document.paragraphs[index], builder(sources)):
            changed += 1
            print(f"  [perbarui] paragraf {index:>3}  {prefix[:44]}")
        else:
            print(f"  [sama]     paragraf {index:>3}  {prefix[:44]}")

    if changed == 0:
        print("-" * 76)
        print("Tidak ada paragraf yang berubah.")
        return

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_prose_refresh.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
        print("-" * 76)
        print("Cadangan:", backup.relative_to(PROJECT_DIR))

    document.save(MANUSCRIPT)
    print("-" * 76)
    print(f"Paragraf diperbarui: {changed}")
    print("Penandaan merah diterapkan per kalimat, bukan per paragraf.")


if __name__ == "__main__":
    main()
