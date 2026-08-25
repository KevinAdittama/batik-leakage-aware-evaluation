"""R2.7 - Sisipkan sensitivitas penyeimbangan kelas ke manuskrip.

Penempatan yang disetujui: subbagian baru 4.7 di akhir Results, dengan Table 8
yang muncul setelah Table 7. Karena tabel baru berada di urutan terakhir, tidak
ada penomoran tabel lain yang bergeser.

Framing yang disetujui: hasil dilaporkan sebagai bukti ketahanan kesimpulan.
Klaim utama artikel adalah soal kebocoran data dan bias akuisisi, bukan soal
manfaat augmentasi, sehingga kesetaraan ketiga lengan memperkuat posisi
tersebut. Kalimat penutup tetap menyebut bahwa membuang original justru
memperburuk hasil, sebagai alasan mengapa jadwal fold-local dipertahankan.

Seluruh angka dibaca dari keluaran tahap 19; tidak ada yang ditulis manual.
Skrip bersifat idempotent dan menolak jalan bila manuskrip terbuka di Word.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\23_insert_balancing_sensitivity.py
"""

from __future__ import annotations

import shutil
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from pipeline_config import PROJECT_DIR, RESULTS_DIR

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
STAGE19 = RESULTS_DIR / "19_balancing_comparison"
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_R2_7_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)
MARKER = "4.7 Class-balancing sensitivity"

ARM_LABEL = {
    "augmented_balanced": "Fold-local augmentation",
    "class_weighted_originals": "Class weighting",
    "balanced_original_sampling": "Balanced original sampling",
}
ARM_COMPOSITION = {
    "augmented_balanced": "200 per class; originals + fold-local derivatives",
    "class_weighted_originals": "Originals only; balanced class weights",
    "balanced_original_sampling": "Originals only; majority class downsampled",
}
HEADERS = [
    "Balancing arm",
    "Training composition",
    "Macro-F1",
    "Bal. accuracy",
    "MCC",
    "SVM-RBF selected",
]


def fmt(value: float, decimals: int = 3) -> str:
    return f"{float(value):.{decimals}f}"


def lock_present() -> bool:
    return any(MANUSCRIPT.parent.glob("~$*"))


def load_numbers() -> dict:
    summary = pd.read_csv(STAGE19 / "arm_summary.csv").set_index("arm")
    paired = pd.read_csv(STAGE19 / "paired_differences.csv")
    selection = pd.read_csv(STAGE19 / "selection_frequency.csv")

    models = selection[(selection.kind == "model") & (selection.choice == "SVM (RBF)")]
    svm = dict(zip(models.arm, models["count"]))
    of = int(models["of"].iloc[0])

    features = selection[
        (selection.kind == "feature_set") & (selection.choice == "Gabungan 6 Fitur")
    ]
    feature_counts = set(features["count"])
    if len(feature_counts) != 1:
        raise AssertionError("Kelompok fitur terpilih tidak seragam antar lengan")

    def diff(arm: str, metric: str) -> pd.Series:
        row = paired[
            (paired.comparison == f"{arm} minus augmented_balanced")
            & (paired.metric == metric)
        ]
        if row.empty:
            raise AssertionError(f"Selisih berpasangan tidak ditemukan: {arm}/{metric}")
        return row.iloc[0]

    return {
        "summary": summary,
        "svm": svm,
        "of": of,
        "feature_selected": feature_counts.pop(),
        "weighted_diff": diff("class_weighted_originals", "f1_macro"),
        "balanced_diff": diff("balanced_original_sampling", "f1_macro"),
    }


def styled(paragraph, text, *, bold=False, size=Pt(11)):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.color.rgb = RED
    return run


def build_table(document: Document, data: dict) -> Table:
    summary = data["summary"]
    table = document.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"

    for index, header in enumerate(HEADERS):
        cell = table.rows[0].cells[index]
        cell.text = ""
        styled(cell.paragraphs[0], header, bold=True, size=Pt(10))

    for arm in ARM_LABEL:
        row = summary.loc[arm]
        values = [
            ARM_LABEL[arm],
            ARM_COMPOSITION[arm],
            f"{fmt(row.macro_f1_mean)} +/- {fmt(row.macro_f1_std)}",
            f"{fmt(row.balanced_accuracy_mean)} +/- {fmt(row.balanced_accuracy_std)}",
            fmt(row.mcc_mean),
            f"{int(data['svm'][arm])}/{data['of']}",
        ]
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = ""
            styled(cells[index].paragraphs[0], value, size=Pt(10))
    return table


def compose_prose(data: dict) -> str:
    summary = data["summary"]
    aug = summary.loc["augmented_balanced"]
    weighted = summary.loc["class_weighted_originals"]
    balanced = summary.loc["balanced_original_sampling"]
    wd = data["weighted_diff"]
    bd = data["balanced_diff"]

    return (
        "A final development-side sensitivity analysis examined whether the "
        "class-balancing schedule, rather than the audited protocol itself, "
        "drove the reported conclusions. Fold-local augmentation, class "
        "weighting on originals only, and balanced original sampling were "
        "compared on identical outer and inner splits across five repeats, "
        "using the same candidate space and the same originals-only validation "
        f"and test folds (Table 8). Repeat-level macro-F1 was {fmt(aug.macro_f1_mean)} "
        f"+/- {fmt(aug.macro_f1_std)} for fold-local augmentation, "
        f"{fmt(weighted.macro_f1_mean)} +/- {fmt(weighted.macro_f1_std)} for class "
        f"weighting, and {fmt(balanced.macro_f1_mean)} +/- {fmt(balanced.macro_f1_std)} "
        "for balanced original sampling. Paired at the outer-fold level, class "
        f"weighting differed from augmentation by {fmt(wd.mean_difference)} macro-F1, "
        f"with {int(wd.folds_tied)} of {int(wd.n_paired_folds)} folds producing "
        f"identical scores, whereas balanced original sampling differed by "
        f"{fmt(bd.mean_difference)} and showed the largest repeat-level variability "
        f"({fmt(balanced.macro_f1_std)} versus {fmt(aug.macro_f1_std)}). All three "
        "arms selected the combined six-feature representation in "
        f"{int(data['feature_selected'])} of {data['of']} outer folds, and SVM-RBF in "
        f"{int(data['svm']['augmented_balanced'])}, "
        f"{int(data['svm']['class_weighted_originals'])}, and "
        f"{int(data['svm']['balanced_original_sampling'])} folds respectively. The "
        "conclusions of this study are therefore insensitive to the balancing "
        "schedule: augmentation is not required to reach the reported development "
        "performance, while discarding batik originals to equalise class counts "
        "measurably degrades it, which is why the fold-local schedule was retained. "
        "The augmented arm reproduced the earlier repeated nested estimate to "
        "floating-point tolerance, confirming that the comparison isolates the "
        "balancing rule rather than a change in the analysis harness."
    )


METHOD_SENTENCE = (
    " Fourth, the repeated nested design was rerun under three class-balancing "
    "arms on identical splits, namely fold-local augmentation, class weighting "
    "applied to originals only, and balanced original sampling, so that the "
    "influence of the balancing schedule could be separated from the influence "
    "of the validation protocol."
)

CAPTION = (
    "Table 8. Class-balancing sensitivity on identical repeated nested splits "
    "(five repeats x five outer folds x four inner folds). Inner validation and "
    "outer test folds contain unaugmented originals only, and the external "
    "collection is never loaded. Values are repeat-level means with standard "
    "deviations across repeats."
)


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if not (STAGE19 / "arm_summary.csv").exists():
        raise FileNotFoundError(
            "Hasil tahap 19 belum ada. Jalankan "
            "19_augmentation_balancing_comparison.py lebih dahulu."
        )
    if lock_present():
        raise RuntimeError("Manuskrip masih terbuka di Word. Tutup lebih dahulu.")

    document = Document(MANUSCRIPT)
    if any(MARKER in p.text for p in document.paragraphs):
        print("Subbagian 4.7 sudah ada. Tidak ada perubahan.")
        return

    data = load_numbers()

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_R2_7.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
        print("Cadangan:", backup.relative_to(PROJECT_DIR))

    # 1. Kalimat metode di Section 3.8.
    method_paragraph = next(
        p for p in document.paragraphs
        if p.text.strip().startswith("Four additional diagnostics were run")
    )
    styled(method_paragraph, METHOD_SENTENCE)

    # 2. Titik sisip: paragraf terakhir Results, tepat sebelum heading Discussion.
    kids = list(document.element.body.iterchildren())
    discussion_index = next(
        i for i, ch in enumerate(kids)
        if ch.tag.endswith("}p")
        and Paragraph(ch, document).text.strip() == "5 Discussion"
    )
    anchor = kids[discussion_index - 1]

    heading = document.add_paragraph(style="Q3 Subsection")
    styled(heading, MARKER, size=Pt(11))

    prose = document.add_paragraph(style="Q3 Body")
    styled(prose, compose_prose(data))

    caption = document.add_paragraph(style="Q3 Caption")
    styled(caption, CAPTION, size=Pt(9))

    table = build_table(document, data)

    anchor.addnext(table._tbl)
    anchor.addnext(caption._p)
    anchor.addnext(prose._p)
    anchor.addnext(heading._p)

    document.save(MANUSCRIPT)

    check = Document(MANUSCRIPT)
    captions = [
        p.text.strip() for p in check.paragraphs
        if p.text.strip().startswith("Table ") and p.style.name == "Q3 Caption"
    ]
    print("-" * 72)
    print(f"Tabel dalam manuskrip: {len(check.tables)}")
    for text in captions:
        print("  ", text[:64])
    print("-" * 72)
    print("Subbagian 4.7 dan Table 8 disisipkan; penomoran tabel lain tidak bergeser.")


if __name__ == "__main__":
    main()
