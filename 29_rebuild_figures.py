"""Bangun ulang Figure 6, 7, dan 8, lalu tukar gambar tertanam di manuskrip.

Latar belakang
--------------
Setelah fold dipisahkan pada grain foto sumber, model formal berpindah dari
Random Forest ke SVM-RBF dan seluruh angka sisi development berubah. Tiga gambar
karena itu menjadi basi:

  Figure 6 - macro-F1 lima fold per kelompok fitur dan classifier.
  Figure 7 - confusion matrix eksternal model formal dan ResNet18.
  Figure 8 - analisis kesalahan eksternal model formal.

Gaya visualnya ditiru dari gambar yang sedang terpasang: serif, colormap biru
untuk confusion matrix, hijau untuk batik dan biru untuk non-batik pada batang
subjenis, serta anotasi merah pada contoh kesalahan.

Label (a) dan (b) tidak lagi digambar di dalam panel. Editor butir E.5 melarang
menyisipkan (a), (b), (c) ke dalam gambar, sedangkan versi lama memuatnya. Panel
kini dibedakan oleh judulnya sendiri, dan caption merujuk kiri atau kanan.

Prinsip
-------
1. Tidak ada nama model yang ditulis harfiah. Model formal dibaca dari
   `selected_model_result.json`, sama seperti tahap 13.
2. Bagian gambar di DOCX ditukar isinya, bukan disisipkan sebagai gambar baru.
   Menambah gambar baru akan menggeser tata letak; menukar isi bagian yang sudah
   ada mempertahankan ukuran dan posisi.
3. Caption yang menyebut kondisi lama ikut diperbarui dan ditandai merah.
4. Idempotent, dan menolak jalan bila manuskrip terbuka di Word.

Setelah dijalankan, tata letak wajib diperiksa visual di Word. Skrip ini menjaga
ukuran bingkai, bukan menjamin pemenggalan halaman tetap sama.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\29_rebuild_figures.py
"""

from __future__ import annotations

import json
import shutil
from datetime import date
from pathlib import Path

import matplotlib
import numpy as np
import pandas as pd

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import RGBColor  # noqa: E402

from pipeline_common import read_image_color, resolve_project_path  # noqa: E402
from pipeline_config import PROJECT_DIR, RESULTS_DIR  # noqa: E402

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
AUDIT_OUT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "04_Tabel_Manifest_dan_Hasil"
    / "Audit_Numerik_dan_Eksperimen"
    / "outputs"
)
FIGURE_DIR = RESULTS_DIR / "29_figures"
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_Gambar_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)
BATIK_GREEN = "#348B5E"
NON_BATIK_BLUE = "#3E75AF"

plt.rcParams.update({
    "font.family": "serif",
    "font.serif": ["DejaVu Serif"],
    "axes.titlesize": 12,
    "axes.labelsize": 10,
    "xtick.labelsize": 9,
    "ytick.labelsize": 9,
})


def lock_present() -> bool:
    return any(MANUSCRIPT.parent.glob("~$*"))


def formal_model() -> tuple[str, str]:
    selected = json.loads(
        (RESULTS_DIR / "08_uji_eksternal/selected_model_result.json").read_text(
            encoding="utf-8"
        )
    )
    return str(selected["model"]), str(selected["model_slug"])


def confusion_panel(axis, matrix: np.ndarray, title: str) -> None:
    axis.imshow(matrix, cmap="Blues", vmin=0, vmax=matrix.max())
    labels = ["Non-batik", "Batik"]
    axis.set_xticks([0, 1], labels)
    axis.set_yticks([0, 1], labels)
    axis.set_xlabel("Predicted class")
    axis.set_ylabel("True class")
    axis.set_title(title, fontweight="bold")
    threshold = matrix.max() * 0.6
    for row in range(2):
        for column in range(2):
            value = int(matrix[row, column])
            axis.text(
                column, row, str(value), ha="center", va="center",
                fontsize=20, fontweight="bold",
                color="white" if value > threshold else "#1F2937",
            )
    for spine in axis.spines.values():
        spine.set_visible(False)
    axis.tick_params(length=3)


def build_figure7(output: Path) -> dict:
    """Confusion matrix eksternal: model formal dan ResNet18."""
    name, slug = formal_model()
    audited = pd.read_csv(AUDIT_OUT / "external_confusion_matrices_audited.csv")
    metrics = pd.read_csv(AUDIT_OUT / "external_metrics_full_precision.csv").set_index("model")

    def matrix_for(model: str) -> np.ndarray:
        rows = audited[audited["model"] == model]
        return (
            rows.pivot(index="true_label", columns="predicted_label", values="n")
            .sort_index().sort_index(axis=1).to_numpy()
        )

    audit_name = name.replace("SVM (RBF)", "SVM-RBF")
    figure, axes = plt.subplots(1, 2, figsize=(10.2, 4.4))
    confusion_panel(
        axes[0], matrix_for(audit_name),
        f"{audit_name}\nmacro-F1 = {metrics.loc[audit_name, 'macro_f1']:.3f}",
    )
    confusion_panel(
        axes[1], matrix_for("ResNet18"),
        f"ResNet18 (post-hoc)\nmacro-F1 = {metrics.loc['ResNet18', 'macro_f1']:.3f}",
    )
    figure.tight_layout()
    figure.savefig(output, dpi=300, bbox_inches="tight")
    plt.close(figure)
    return {"model": audit_name, "slug": slug}


def error_examples(limit: int = 2) -> list[dict]:
    """Kesalahan berskor paling ekstrem pada kedua arah, dari model formal."""
    _, slug = formal_model()
    predictions = pd.read_csv(RESULTS_DIR / f"08_uji_eksternal/predictions_{slug}.csv")
    wrong = predictions[predictions.label != predictions.predicted_label]
    false_negative = wrong[wrong.label == 1].nsmallest(limit, "score_batik")
    false_positive = wrong[wrong.label == 0].nlargest(limit, "score_batik")
    chosen = []
    for record in pd.concat([false_negative, false_positive]).to_dict("records"):
        direction = (
            "Batik -> Non-batik" if record["label"] == 1 else "Non-batik -> Batik"
        )
        chosen.append({
            "path": record["path"],
            "caption": (
                f"{direction}\n{str(record['subjenis']).replace('_', ' ')}; "
                f"p(batik)={record['score_batik']:.2f}"
            ),
        })
    return chosen


def build_figure8(output: Path) -> dict:
    """(a) tingkat kesalahan per subjenis, (b) contoh kesalahan berskor ekstrem."""
    name, slug = formal_model()
    audit_name = name.replace("SVM (RBF)", "SVM-RBF")
    subtype = pd.read_csv(AUDIT_OUT / f"{slug}_subtype_error_counts_audited.csv")
    affected = subtype[subtype.errors > 0].sort_values("error_rate")
    affected = affected.assign(
        display=lambda frame: frame.apply(
            lambda row: (
                f"{'B' if row['kelas'] == 'batik' else 'NB'}: "
                f"{str(row['subjenis']).replace('_', ' ')}"
            ),
            axis=1,
        )
    )

    figure = plt.figure(figsize=(13.6, 6.4))
    grid = figure.add_gridspec(2, 4, width_ratios=[2.2, 0.12, 1, 1], hspace=0.42, wspace=0.18)

    axis = figure.add_subplot(grid[:, 0])
    colors = [
        BATIK_GREEN if value == "batik" else NON_BATIK_BLUE
        for value in affected["kelas"]
    ]
    bars = axis.barh(affected["display"], affected["error_rate"], color=colors)
    axis.bar_label(
        bars,
        labels=[
            f"{int(errors)}/{int(total)}"
            for errors, total in zip(affected["errors"], affected["n"])
        ],
        padding=4, fontsize=9,
    )
    axis.set_xlim(0, 1.12)
    axis.set_xlabel(f"External error rate ({audit_name})")
    # Judul kedua panel ditempatkan sebagai teks figure agar sejajar; judul axes
    # akan mengikuti tinggi axes-nya sendiri dan terlihat turun sebelah.
    figure.text(0.30, 0.965, "Descriptive subtype errors",
                ha="center", fontsize=12, fontweight="bold")
    axis.grid(axis="x", alpha=0.2, linewidth=0.5)
    for spine in ("top", "right"):
        axis.spines[spine].set_visible(False)

    examples = error_examples()
    positions = [(0, 2), (0, 3), (1, 2), (1, 3)]
    for (row, column), example in zip(positions, examples):
        panel = figure.add_subplot(grid[row, column])
        image = read_image_color(resolve_project_path(example["path"]))
        panel.imshow(image[:, :, ::-1])
        panel.set_title(example["caption"], fontsize=9, color="#B03A48")
        panel.axis("off")
    if examples:
        figure.text(
            0.685, 0.965, "Predicted-score errors",
            ha="center", fontsize=12, fontweight="bold",
        )

    figure.savefig(output, dpi=300, bbox_inches="tight")
    plt.close(figure)
    return {
        "model": audit_name,
        "errors": int(subtype.errors.sum()),
        "false_negative": int(subtype.false_negative.sum()),
        "false_positive": int(subtype.false_positive.sum()),
    }


def replace_image(document: Document, caption_prefix: str, source: Path) -> bool:
    """Tukar isi bagian gambar milik caption tertentu."""
    from docx.text.paragraph import Paragraph

    blip_tag = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    embed_attr = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )

    last_rid = None
    for child in document.element.body.iterchildren():
        if not child.tag.endswith("}p"):
            continue
        blips = child.findall(f".//{blip_tag}")
        if blips:
            last_rid = blips[0].get(embed_attr)
        text = Paragraph(child, document).text.strip()
        if text.startswith(caption_prefix):
            if last_rid is None:
                raise AssertionError(f"Tidak ada gambar sebelum caption {caption_prefix!r}")
            part = document.part.related_parts[last_rid]
            blob = source.read_bytes()
            if part._blob == blob:
                return False
            part._blob = blob
            return True
    raise AssertionError(f"Caption tidak ditemukan: {caption_prefix!r}")


def retitle(document: Document, prefix: str, new_text: str) -> bool:
    """Ganti caption dan tandai merah, karena isinya berubah."""
    for paragraph in document.paragraphs:
        if not paragraph.text.strip().startswith(prefix):
            continue
        if paragraph.text.strip() == new_text:
            return False
        template = paragraph.runs[0] if paragraph.runs else None
        size = template.font.size if template is not None else None
        for run in list(paragraph.runs):
            run._r.getparent().remove(run._r)
        run = paragraph.add_run(new_text)
        if size is not None:
            run.font.size = size
        run.font.color.rgb = RED
        return True
    raise AssertionError(f"Caption tidak ditemukan: {prefix!r}")


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if lock_present():
        raise RuntimeError("Manuskrip masih terbuka di Word. Tutup lebih dahulu.")

    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    figure7 = FIGURE_DIR / "figure7_external_confusion.png"
    figure8 = FIGURE_DIR / "figure8_error_analysis.png"

    print("=" * 76)
    print("BANGUN ULANG GAMBAR DAN TUKAR KE MANUSKRIP")
    print("=" * 76)
    info7 = build_figure7(figure7)
    info8 = build_figure8(figure8)
    print(f"  Figure 7 dibangun untuk {info7['model']}")
    print(f"  Figure 8 dibangun untuk {info8['model']}: {info8['errors']} error "
          f"({info8['false_negative']} FN + {info8['false_positive']} FP)")

    document = Document(MANUSCRIPT)
    model = info7["model"]

    swaps = [
        ("Figure 6.", RESULTS_DIR / "07_cross_validation/feature_ablation_cv.png"),
        ("Figure 7.", figure7),
        ("Figure 8.", figure8),
    ]
    captions = [
        (
            "Figure 1.",
            "Figure 1. Source-group-aware leakage-aware experimental design. Splits "
            "are formed over source-photograph groups before augmentation; "
            "validation and external images remain unaugmented. The frozen deep "
            "branch is a post-hoc benchmark.",
        ),
        (
            "Figure 7.",
            f"Figure 7. External confusion matrices for the handcrafted {model} "
            "returned by the CV selection rule, left, and the post-hoc ResNet18 "
            "frozen-feature benchmark, right. The selection margin between "
            "handcrafted models is smaller than the fold-level standard error "
            "(Table 10), so the left matrix is not a claim of superiority.",
        ),
        (
            "Figure 8.",
            f"Figure 8. {model} external error analysis. In the left panel, B and NB "
            "denote batik and non-batik, and labels show errors over totals. In the "
            "right panel, prediction scores were not probability-calibrated.",
        ),
    ]

    changed = []
    for prefix, source in swaps:
        if not source.exists():
            raise FileNotFoundError(source)
        if replace_image(document, prefix, source):
            changed.append(f"gambar {prefix.rstrip('.')}")
    for prefix, text in captions:
        if retitle(document, prefix, text):
            changed.append(f"caption {prefix.rstrip('.')}")

    if not changed:
        print("-" * 76)
        print("Tidak ada perubahan. Gambar dan caption sudah selaras.")
        return

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_figure_rebuild.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
        print("-" * 76)
        print("Cadangan:", backup.relative_to(PROJECT_DIR))

    document.save(MANUSCRIPT)
    print("-" * 76)
    for item in changed:
        print(f"  diperbarui: {item}")
    print("-" * 76)
    print("Ukuran bingkai gambar dipertahankan, tetapi pemenggalan halaman")
    print("harus diperiksa visual di Word sebelum submit.")


if __name__ == "__main__":
    main()
