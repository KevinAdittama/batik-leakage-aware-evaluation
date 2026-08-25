"""R1.9 - Tabel komposisi fold pelatihan per kelas.

Reviewer 1 butir 9 meminta laporan jumlah original, derivative, dan total
instance pelatihan untuk setiap fold dan setiap kelas, beserta konfirmasi bahwa
tidak ada original yang menyumbang derivative secara tidak proporsional.

Skrip ini hanya membaca sumber terauditkan yang sudah ada
(fold_composition_summary.csv, hasil 13_numerical_audit.py) dan mengubahnya
menjadi bahan siap pakai untuk manuskrip. Skrip ini tidak menghitung ulang
fold, tidak melatih model, dan tidak menyentuh manuskrip.

Keluaran:
  - fold_composition_table_manuscript.csv : tabel siap tampil
  - Tabel_Komposisi_Fold_R1_9.docx        : satu tabel Word siap disalin
  - fold_composition_prose_draft.txt      : draf kalimat pengiring
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from pipeline_config import PROJECT_DIR

AUDIT_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "04_Tabel_Manifest_dan_Hasil"
    / "Audit_Numerik_dan_Eksperimen"
    / "outputs"
)
SOURCE = AUDIT_DIR / "fold_composition_summary.csv"
OUT_DIR = AUDIT_DIR / "R1_9_fold_composition"

RED = RGBColor(0xFF, 0x00, 0x00)
KELAS_TAMPIL = {"batik": "Batik", "non_batik": "Non-batik"}

HEADERS = [
    "Fold",
    "Class",
    "Val. originals",
    "Train originals",
    "Train derivatives",
    "Train total",
    "Derivatives per original (min-max)",
    "Mean derivatives per original",
]


def build_table(source: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, r in source.iterrows():
        rows.append(
            {
                "Fold": int(r["fold"]),
                "Class": KELAS_TAMPIL[r["kelas"]],
                "Val. originals": int(r["validation_original"]),
                "Train originals": int(r["train_original"]),
                "Train derivatives": int(r["train_derivative"]),
                "Train total": int(r["train_total"]),
                "Derivatives per original (min-max)": (
                    f"{int(r['derivatives_per_original_min'])}-"
                    f"{int(r['derivatives_per_original_max'])}"
                ),
                "Mean derivatives per original": (
                    f"{float(r['derivatives_per_original_mean']):.2f}"
                ),
            }
        )
    return pd.DataFrame(rows, columns=HEADERS)


def verify(source: pd.DataFrame) -> dict:
    """Assertion yang menegakkan klaim yang akan ditulis di manuskrip."""
    checks = {}

    totals = source["train_total"].unique().tolist()
    assert totals == [200], f"train_total tidak seragam 200: {totals}"
    checks["setiap fold/kelas berisi 200 training instance"] = True

    sums = (source["train_original"] + source["train_derivative"]) - source["train_total"]
    assert (sums == 0).all(), "original + derivative tidak sama dengan total"
    checks["original + derivative = total di semua baris"] = True

    per_fold_val = source.groupby("fold")["validation_original"].sum()
    assert per_fold_val.tolist() == [41, 40, 40, 40, 40], per_fold_val.tolist()
    checks["ukuran fold validasi 41/40/40/40/40"] = True

    assert int(per_fold_val.sum()) == 201, "total original bukan 201"
    checks["total original development 201"] = True

    batik = source[source["kelas"] == "batik"]
    non_batik = source[source["kelas"] == "non_batik"]
    assert batik["derivatives_per_original_max"].max() == 2
    assert non_batik["derivatives_per_original_max"].max() == 4
    checks["derivative maksimum per original: batik 2, non-batik 4"] = True

    # Konfirmasi tidak ada original yang menyumbang derivative tak proporsional:
    # batas atas per kelas dipatuhi di seluruh fold.
    for _, r in source.iterrows():
        dist = json.loads(r["derivative_count_distribution"])
        n_sources = sum(dist.values())
        assert n_sources == int(r["train_original"]), (
            f"distribusi derivative fold {r['fold']} {r['kelas']} "
            f"tidak mencakup semua original"
        )
        weighted = sum(int(k) * v for k, v in dist.items())
        assert weighted == int(r["train_derivative"]), (
            f"jumlah derivative fold {r['fold']} {r['kelas']} tidak konsisten"
        )
    checks["distribusi derivative konsisten dengan jumlah original & derivative"] = True

    return checks


def write_docx(table: pd.DataFrame, path: Path) -> None:
    doc = Document()
    doc.core_properties.title = "Tabel komposisi fold pelatihan (R1.9)"

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = caption.add_run(
        "Table X. Training-fold composition per class after the approved two-image "
        "exclusion. Every training split contains 200 instances per class; "
        "derivatives are training-only descendants of originals in the same "
        "training fold. Validation folds contain unaugmented originals only."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RED

    word_table = doc.add_table(rows=1, cols=len(HEADERS))
    word_table.style = "Table Grid"
    for i, h in enumerate(HEADERS):
        cell = word_table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RED

    for _, row in table.iterrows():
        cells = word_table.add_row().cells
        for i, h in enumerate(HEADERS):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(row[h]))
            r.font.size = Pt(9)
            r.font.color.rgb = RED

    doc.save(path)


PROSE = (
    "Training-fold composition is reported per fold and per class in Table X. "
    "The five validation folds contained 41, 40, 40, 40, and 40 unaugmented "
    "originals, together covering all 201 clean development files exactly once. "
    "Within each training split, batik contributed 109-110 originals and 90-91 "
    "derivatives, while non-batik contributed 51-52 originals and 148-149 "
    "derivatives, so every training split held 200 instances per class. "
    "Because the two classes differ in the number of available originals, the "
    "balancing schedule drew at most two derivatives per batik original and at "
    "most four per non-batik original; no original exceeded these per-class "
    "limits in any fold. The resulting class asymmetry in derivative density is "
    "a property of the balancing schedule rather than of any individual source, "
    "and its effect on model selection is examined in the repeated strictly "
    "nested sensitivity analysis."
)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    source = pd.read_csv(SOURCE)

    checks = verify(source)
    table = build_table(source)

    csv_path = OUT_DIR / "fold_composition_table_manuscript.csv"
    docx_path = OUT_DIR / "Tabel_Komposisi_Fold_R1_9.docx"
    prose_path = OUT_DIR / "fold_composition_prose_draft.txt"

    table.to_csv(csv_path, index=False)
    write_docx(table, docx_path)
    prose_path.write_text(PROSE + "\n", encoding="utf-8")

    print("Sumber :", SOURCE.relative_to(PROJECT_DIR))
    print("Keluaran:", OUT_DIR.relative_to(PROJECT_DIR))
    print("-" * 70)
    print(table.to_string(index=False))
    print("-" * 70)
    for label in checks:
        print(f"  [OK] {label}")
    print("-" * 70)
    for p in (csv_path, docx_path, prose_path):
        print("  tersimpan:", p.name)


if __name__ == "__main__":
    main()
