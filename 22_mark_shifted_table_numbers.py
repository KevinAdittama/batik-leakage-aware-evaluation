"""E.1 - Tandai merah angka tabel yang bergeser akibat penyisipan Table 5.

Penyisipan tabel komposisi fold (butir R1.9) menggeser dua caption:

    Table 5. Five-fold results ...      -> Table 6. Five-fold results ...
    Table 6. Classical and frozen-deep  -> Table 7. Classical and frozen-deep

Isi caption tidak berubah, hanya nomornya. Editor butir E.1 meminta setiap
modifikasi ditandai merah, tetapi memerahkan seluruh caption akan menyesatkan
karena memberi kesan tabelnya baru atau isinya berubah. Karena itu hanya
angkanya yang diberi warna merah.

Skrip bersifat idempotent: bila angka sudah merah, tidak ada perubahan.
Manuskrip harus tertutup di Word sebelum dijalankan.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\22_mark_shifted_table_numbers.py
"""

from __future__ import annotations

import copy
import re
import shutil
from datetime import date

from docx import Document
from docx.shared import RGBColor

from pipeline_config import PROJECT_DIR

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_Penandaan_Nomor_Tabel_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)

# Caption yang nomornya bergeser, beserta angka yang harus dimerahkan.
TARGETS = [
    ("Table 6. Five-fold results", "6"),
    ("Table 7. Classical and frozen-deep", "7"),
]


def lock_file_present() -> bool:
    return any(MANUSCRIPT.parent.glob(f"~${MANUSCRIPT.stem[:20]}*"))


def split_run(paragraph, run, numeral: str) -> bool:
    """Pecah run sehingga angka nomor tabel berdiri sendiri, lalu warnai merah.

    Mengembalikan True bila terjadi perubahan.
    """
    text = run.text
    match = re.search(rf"(?<=Table ){re.escape(numeral)}(?=\.)", text)
    if not match:
        return False

    before, numeral_text, after = (
        text[: match.start()],
        match.group(0),
        text[match.end():],
    )

    run.text = before
    element = run._r
    parent = element.getparent()
    index = parent.index(element)

    numeral_run = copy.deepcopy(element)
    parent.insert(index + 1, numeral_run)

    tail_run = copy.deepcopy(element)
    parent.insert(index + 2, tail_run)

    from docx.text.run import Run

    numeral_obj = Run(numeral_run, paragraph)
    numeral_obj.text = numeral_text
    numeral_obj.font.color.rgb = RED

    Run(tail_run, paragraph).text = after
    return True


def already_marked(paragraph, numeral: str) -> bool:
    for run in paragraph.runs:
        if run.text.strip() != numeral:
            continue
        colour = run.font.color
        if colour is not None and colour.rgb == RED:
            return True
    return False


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if lock_file_present():
        raise RuntimeError(
            "Manuskrip masih terbuka di Word. Tutup dokumennya lebih dahulu."
        )

    document = Document(MANUSCRIPT)

    pending = []
    for prefix, numeral in TARGETS:
        paragraph = next(
            (p for p in document.paragraphs if p.text.strip().startswith(prefix)),
            None,
        )
        if paragraph is None:
            raise RuntimeError(f"Caption tidak ditemukan: {prefix}")
        if already_marked(paragraph, numeral):
            print(f"[lewati] '{prefix}' sudah merah pada angkanya")
            continue
        pending.append((paragraph, prefix, numeral))

    if not pending:
        print("Tidak ada perubahan; seluruh nomor tabel yang bergeser sudah ditandai.")
        return

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_numeral_marking.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
        print("Cadangan:", backup.relative_to(PROJECT_DIR))

    for paragraph, prefix, numeral in pending:
        changed = any(
            split_run(paragraph, run, numeral) for run in list(paragraph.runs)
        )
        if not changed:
            raise RuntimeError(f"Gagal memisahkan angka pada caption: {prefix}")
        print(f"[tandai] '{prefix}' -> angka {numeral} kini merah")

    document.save(MANUSCRIPT)

    check = Document(MANUSCRIPT)
    for prefix, numeral in TARGETS:
        paragraph = next(
            p for p in check.paragraphs if p.text.strip().startswith(prefix)
        )
        if not already_marked(paragraph, numeral):
            raise RuntimeError(f"Verifikasi gagal untuk: {prefix}")
        colours = {
            str(run.font.color.rgb)
            for run in paragraph.runs
            if run.text.strip() and run.font.color and run.font.color.rgb
        }
        print(f"[verifikasi] {prefix[:34]:<36} warna dalam caption: {colours}")

    print("-" * 70)
    print("Selesai. Hanya angka nomor tabel yang merah; teks caption tetap hitam.")


if __name__ == "__main__":
    main()
