"""Sisipkan subbagian 4.8 beserta Table 9 dan Table 10 ke manuskrip.

Table 9  - adjudikasi grup foto sumber (tahap 24).
Table 10 - stabilitas seleksi model formal (tahap 25).

Penempatan
----------
Keduanya diletakkan di akhir Results, setelah Table 8. Karena berada paling
belakang, tidak ada penomoran tabel lain yang bergeser. Ini aturan yang sama
dengan penyisipan Table 8 oleh skrip 23.

Framing yang dipegang
---------------------
Table 9 melaporkan cacat yang ditemukan pada data sendiri dan cara
memperbaikinya, bukan klaim bahwa koleksi kini bersih. Batas metodenya
dinyatakan terbuka: penyelarasan hanya mendeteksi tumpang tindih translasional.

Table 10 melaporkan bahwa aturan seleksi formal berhenti memisahkan dua model
teratas. Itu disajikan sebagai temuan tentang keterbatasan protokol pada koleksi
kecil, bukan sebagai perbandingan siapa yang menang.

Seluruh angka dibaca dari keluaran tahap 24 dan 25. Skrip bersifat idempotent
dan menolak jalan bila manuskrip terbuka di Word.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\28_insert_group_and_selection_tables.py
"""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from pipeline_config import PROJECT_DIR, RESULTS_DIR

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_Table9_10_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)
MARKER = "4.8 Source-group adjudication and selection stability"

TABLE9_HEADERS = [
    "Subtype", "Image pair", "Offset (px)", "RGB residual",
    "Alignment score", "Folds before grouping",
]
TABLE10_HEADERS = [
    "Split protocol", "Leading model", "Second model",
    "Margin", "Fold-level SE", "Margin / SE", "Nested selection",
]

TABLE9_CAPTION = (
    "Table 9. Same-photograph pairs confirmed by full-resolution alignment among "
    "all pairs of development and external images. Perceptual hashing is not shift "
    "invariant and flagged only two of these pairs. RGB residual is the mean "
    "absolute difference over the aligned overlap on a 0-255 scale; it is never "
    "zero because each crop was resized and re-encoded separately."
)
TABLE10_CAPTION = (
    "Table 10. Stability of the formal selection rule under file-level and "
    "source-group-aware splitting. The selection rule is unchanged in both rows: "
    "highest mean cross-validated macro-F1 on the combined six features. Margin is "
    "the macro-F1 gap between the two leading handcrafted models, and the fold-level "
    "standard error is the pooled fold-to-fold standard deviation divided by the "
    "square root of the number of folds."
)

METHOD_SENTENCE = (
    " Fifth, every pair of development and external images was screened by "
    "normalized cross-correlation and then verified by full-resolution alignment, "
    "so that overlapping crops of one photograph could be grouped before splitting "
    "rather than treated as independent files."
)


def lock_present() -> bool:
    return any(MANUSCRIPT.parent.glob("~$*"))


def fmt(value: float, decimals: int = 3) -> str:
    return f"{float(value):.{decimals}f}"


def styled(paragraph, text, *, bold=False, size=Pt(11)):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.color.rgb = RED
    return run


def build_table(document: Document, headers: list[str], rows: list[list[str]]) -> Table:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        styled(cell.paragraphs[0], header, bold=True, size=Pt(10))
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = ""
            styled(cells[index].paragraphs[0], value, size=Pt(10))
    return table


def table9_rows() -> list[list[str]]:
    confirmed = pd.read_csv(RESULTS_DIR / "24_source_groups/confirmed_pairs.csv")
    leakage_path = RESULTS_DIR / "24_source_groups/fold_leakage_check.csv"
    leakage = pd.read_csv(leakage_path) if leakage_path.exists() else pd.DataFrame()
    folds = {
        (record["left"], record["right"]): (record["fold_left"], record["fold_right"])
        for record in leakage.to_dict("records")
    } if not leakage.empty else {}

    rows = []
    for record in confirmed.sort_values("residual_rgb").to_dict("records"):
        left, right = Path(record["left"]), Path(record["right"])
        pair_folds = folds.get((record["left"], record["right"]))
        rows.append([
            left.parent.name.replace("_", " "),
            f"{left.name} / {right.name}",
            f"({int(record['dx'])}, {int(record['dy'])})",
            fmt(record["residual_rgb"], 2),
            fmt(record["align_score"], 4),
            f"{int(pair_folds[0])} vs {int(pair_folds[1])}" if pair_folds else "-",
        ])
    return rows


def table10_rows() -> list[list[str]]:
    margin = pd.read_csv(
        RESULTS_DIR / "25_model_selection_stability/selection_margin.csv"
    )
    selection = pd.read_csv(
        RESULTS_DIR / "25_model_selection_stability/nested_selection_frequency.csv"
    )
    nested = selection[selection.protocol.str.contains("tahap 14")]

    labels = {
        "file-level": "File-level (as originally reported)",
        "source-group-aware": "Source-group-aware (this revision)",
    }
    rows = []
    for record in margin.to_dict("records"):
        protocol = record["protocol"]
        if protocol == "source-group-aware" and not nested.empty:
            split = ", ".join(
                f"{item['model']} {int(item['selections'])}/{int(item['of'])}"
                for item in nested.sort_values("selections", ascending=False).to_dict("records")
            )
        else:
            split = "-"
        rows.append([
            labels.get(protocol, protocol),
            f"{record['winner']} ({fmt(record['winner_macro_f1'])})",
            f"{record['runner_up']} ({fmt(record['runner_up_macro_f1'])})",
            fmt(record["margin"], 4),
            fmt(record["standard_error_of_mean"], 4),
            f"{float(record['margin_in_standard_errors']):.2f}",
            split,
        ])
    return rows


def compose_prose() -> str:
    confirmed = pd.read_csv(RESULTS_DIR / "24_source_groups/confirmed_pairs.csv")
    groups = pd.read_csv(RESULTS_DIR / "24_source_groups/source_groups.csv")
    development = groups[groups["set"] == "development"]
    multi = development[development.group_size > 1]
    margin = pd.read_csv(
        RESULTS_DIR / "25_model_selection_stability/selection_margin.csv"
    ).set_index("protocol")
    current = margin.loc["source-group-aware"]
    before = margin.loc["file-level"]

    return (
        "Two diagnostics were added after the file-level analysis had been "
        "completed, because the perceptual-hash screen reported in Table 4 proved "
        "insufficient. Perceptual hashing compares global structure and is not "
        "invariant to translation, so two crops taken from one photograph at "
        "different offsets receive distant hashes. Every pair of development and "
        "external images was therefore rescreened by normalized cross-correlation "
        "and verified by full-resolution alignment in colour, which confirmed "
        f"{len(confirmed)} same-photograph pairs covering {len(multi)} images in "
        f"{int(multi.group_id.nunique())} source groups (Table 9). Only two of them "
        "had appeared among the perceptual-hash candidates, and all of them had been "
        "assigned to different validation folds, so crops of one photograph were "
        "present on both sides of a split. The development manifest was regrouped "
        f"from {len(development)} independent files to "
        f"{int(development.group_id.nunique())} source groups and every "
        "development-side analysis in this article was recomputed under "
        "source-group-aware splitting. This test detects translational overlap only; "
        "non-overlapping crops, rotated copies, and heavily rescaled copies remain "
        "undetectable, so the grouping is a lower bound rather than a guarantee of "
        "independence. The regrouping also changed which model the formal selection "
        f"rule returns (Table 10). Under file-level splitting the leading margin was "
        f"{fmt(before.margin, 4)} macro-F1, or "
        f"{float(before.margin_in_standard_errors):.2f} fold-level standard errors; "
        f"under source-group-aware splitting it was {fmt(current.margin, 4)}, or "
        f"{float(current.margin_in_standard_errors):.2f} standard errors, and the "
        "two leading models exchanged places. Neither margin is large enough to "
        "distinguish the models, and the repeated nested protocol divided its "
        "selections almost evenly between them. The formal winner is therefore "
        "reported as an artefact of the selection rule rather than as evidence of "
        "superiority, and all classical models are reported side by side throughout."
    )


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if lock_present():
        raise RuntimeError("Manuskrip masih terbuka di Word. Tutup lebih dahulu.")

    document = Document(MANUSCRIPT)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        print("Subbagian 4.8 sudah ada. Tidak ada perubahan.")
        return

    rows9, rows10 = table9_rows(), table10_rows()
    if not rows9 or not rows10:
        raise AssertionError("Sumber tahap 24 atau 25 kosong; jalankan keduanya dulu")

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_table9_10.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
        print("Cadangan:", backup.relative_to(PROJECT_DIR))

    # Kalimat metode di Section 3.8, menyusul empat diagnostik yang sudah ada.
    method = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Four additional diagnostics were run")
    )
    styled(method, METHOD_SENTENCE)

    kids = list(document.element.body.iterchildren())
    discussion_index = next(
        index for index, child in enumerate(kids)
        if child.tag.endswith("}p")
        and Paragraph(child, document).text.strip() == "5 Discussion"
    )
    anchor = kids[discussion_index - 1]

    heading = document.add_paragraph(style="Q3 Subsection")
    styled(heading, MARKER, size=Pt(11))

    prose = document.add_paragraph(style="Q3 Body")
    styled(prose, compose_prose())

    caption9 = document.add_paragraph(style="Q3 Caption")
    styled(caption9, TABLE9_CAPTION, size=Pt(9))
    table9 = build_table(document, TABLE9_HEADERS, rows9)

    caption10 = document.add_paragraph(style="Q3 Caption")
    styled(caption10, TABLE10_CAPTION, size=Pt(9))
    table10 = build_table(document, TABLE10_HEADERS, rows10)

    for element in (
        table10._tbl, caption10._p, table9._tbl, caption9._p, prose._p, heading._p
    ):
        anchor.addnext(element)

    document.save(MANUSCRIPT)

    check = Document(MANUSCRIPT)
    captions = [
        paragraph.text.strip() for paragraph in check.paragraphs
        if paragraph.text.strip().startswith("Table ")
        and paragraph.style.name == "Q3 Caption"
    ]
    print("-" * 72)
    print(f"Tabel dalam manuskrip: {len(check.tables)}")
    for text in captions:
        print("  ", text[:66])
    print("-" * 72)
    print("Subbagian 4.8, Table 9, dan Table 10 disisipkan di akhir Results.")
    print("Penomoran tabel lain tidak bergeser.")


if __name__ == "__main__":
    main()
