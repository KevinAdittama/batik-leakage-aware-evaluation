"""Segarkan angka Tabel 4 sampai 8 pada manuskrip dari sumber terauditkan.

Latar belakang
--------------
Fold kini dipisahkan pada grain foto sumber (tahap 24), bukan pada grain berkas.
Seluruh angka sisi development karena itu dihitung ulang. Skrip ini memindahkan
angka baru ke manuskrip tanpa menyentuh kalimat, struktur, maupun penomoran.

Prinsip yang dipegang
---------------------
1. Tidak ada angka yang diketik manual. Setiap sel dibaca dari CSV terauditkan.
2. Sel yang nilainya benar-benar berubah ditandai merah (Editor butir E.1). Sel
   yang nilainya sama dibiarkan apa adanya, termasuk warnanya. Memerahkan sel
   yang tidak berubah akan menyesatkan pembaca tentang cakupan revisi.
3. Idempotent. Menjalankan dua kali tidak menghasilkan perubahan tambahan.
4. Menolak jalan bila manuskrip masih terbuka di Word.

Yang tidak dikerjakan di sini
-----------------------------
Kalimat naratif yang memuat angka, subbagian baru untuk adjudikasi grup sumber,
dan penggantian narasi model tunggal. Semuanya ditangani skrip 27.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\26_refresh_manuscript_numbers.py
"""

from __future__ import annotations

import shutil
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.table import Table

from pipeline_config import PROJECT_DIR, RESULTS_DIR

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
AUDIT_OUT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "04_Tabel_Manifest_dan_Hasil"
    / "Audit_Numerik_dan_Eksperimen"
    / "outputs"
)
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_Segar_Angka_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)
KELAS_TAMPIL = {"batik": "Batik", "non_batik": "Non-batik"}


def lock_present() -> bool:
    return any(MANUSCRIPT.parent.glob("~$*"))


def fmt(value: float, decimals: int = 3) -> str:
    return f"{float(value):.{decimals}f}"


def pm(mean: float, std: float, decimals: int = 3) -> str:
    return f"{fmt(mean, decimals)} +/- {fmt(std, decimals)}"


def set_cell(cell, text: str) -> bool:
    """Tulis ulang isi sel. Merah hanya bila nilainya benar-benar berubah."""
    if cell.text.strip() == text:
        return False

    paragraph = cell.paragraphs[0]
    template = paragraph.runs[0] if paragraph.runs else None
    size = template.font.size if template is not None else None
    bold = template.bold if template is not None else None

    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)

    run = paragraph.add_run(text)
    if size is not None:
        run.font.size = size
    run.bold = bold
    run.font.color.rgb = RED
    return True


def find_table(document: Document, wanted: set[str]) -> Table:
    for table in document.tables:
        header = {cell.text.strip().lower() for cell in table.rows[0].cells}
        if wanted.issubset(header):
            return table
    raise AssertionError(f"Tabel dengan header {sorted(wanted)} tidak ditemukan")


def refresh_table4(document: Document) -> int:
    """Tabel diagnostik: hasil pHash, nested, nested berulang, kontrol metadata."""
    table = find_table(document, {"diagnostic", "protocol", "development result"})
    groups = pd.read_csv(RESULTS_DIR / "24_source_groups/source_groups.csv")
    confirmed = pd.read_csv(RESULTS_DIR / "24_source_groups/confirmed_pairs.csv")
    multi = groups[groups.group_size > 1]

    nested = pd.read_csv(
        RESULTS_DIR / "12_submission_robustness/nested_cv_summary.csv", index_col=0
    )
    repeated = pd.read_csv(
        RESULTS_DIR / "14_repeated_nested_augmentation/summary.csv"
    ).set_index("metric")
    selection = pd.read_csv(
        RESULTS_DIR / "14_repeated_nested_augmentation/selection_frequency.csv"
    )
    metadata = pd.read_csv(
        RESULTS_DIR / "12_submission_robustness/metadata_negative_control_metrics.csv"
    )
    development = metadata[metadata.evaluation == "development_5fold"]
    external = metadata[metadata.evaluation != "development_5fold"].iloc[0]

    total = int(selection.outer_fold_selections.sum())
    selection_text = "; ".join(
        f"{record['selected_model']} {int(record['outer_fold_selections'])}/{total}"
        for record in selection.sort_values(
            "outer_fold_selections", ascending=False
        ).to_dict("records")
    )

    updates = {
        "Perceptual-hash screen": (
            f"{len(confirmed)} same-photograph pairs confirmed by alignment; "
            f"{int(multi.group_size.count())} images in "
            f"{int(multi.group_id.nunique())} source groups"
        ),
        "Nested model/feature selection": (
            f"Macro-F1 {pm(nested.loc['f1_macro', 'mean'], nested.loc['f1_macro', 'std'])}"
        ),
        "Repeated nested fold-local augmentation": (
            f"Macro-F1 "
            f"{pm(repeated.loc['f1_macro', 'repeat_mean'], repeated.loc['f1_macro', 'repeat_std'])}; "
            f"{selection_text}"
        ),
        "Learned metadata control": (
            f"Macro-F1 {pm(development.f1_macro.mean(), development.f1_macro.std(ddof=1))}"
        ),
    }

    changed = 0
    for row in table.rows[1:]:
        label = row.cells[0].text.strip()
        if label in updates:
            changed += set_cell(row.cells[2], updates[label])
        if label == "Learned metadata control":
            changed += set_cell(
                row.cells[3],
                f"Macro-F1 {fmt(external.f1_macro)}; MCC {fmt(external.mcc)}",
            )
    return changed


def refresh_table5(document: Document) -> int:
    """Komposisi fold pelatihan per kelas."""
    table = find_table(document, {"fold", "class", "val. originals", "train total"})
    source = pd.read_csv(AUDIT_OUT / "fold_composition_summary.csv")
    if len(table.rows) - 1 != len(source):
        raise AssertionError("Jumlah baris Table 5 tidak cocok dengan sumber terauditkan")

    changed = 0
    for row, (_, record) in zip(table.rows[1:], source.iterrows()):
        values = [
            str(int(record["fold"])),
            KELAS_TAMPIL[record["kelas"]],
            str(int(record["validation_original"])),
            str(int(record["train_original"])),
            str(int(record["train_derivative"])),
            str(int(record["train_total"])),
            f"{int(record['derivatives_per_original_min'])}-"
            f"{int(record['derivatives_per_original_max'])}",
            f"{float(record['derivatives_per_original_mean']):.2f}",
        ]
        for cell, value in zip(row.cells, values):
            changed += set_cell(cell, value)
    return changed


def refresh_table6(document: Document) -> int:
    """Hasil 5-fold enam fitur gabungan untuk ketiga model klasik."""
    table = find_table(document, {"model", "accuracy", "bal. accuracy", "macro-f1"})
    source = pd.read_csv(
        RESULTS_DIR / "07_cross_validation/cv_summary_primary.csv"
    )
    alias = {"SVM (RBF)": "SVM-RBF"}
    source["display"] = source.model.map(lambda name: alias.get(name, name))
    lookup = source.set_index("display")

    changed = 0
    for row in table.rows[1:]:
        model = row.cells[0].text.strip()
        if model not in lookup.index:
            raise AssertionError(f"Model {model!r} tidak ada pada cv_summary_primary")
        record = lookup.loc[model]
        values = [
            pm(record.accuracy_mean, record.accuracy_std),
            pm(record.balanced_accuracy_mean, record.balanced_accuracy_std),
            pm(record.f1_macro_mean, record.f1_macro_std),
            pm(record.mcc_mean, record.mcc_std),
            pm(record.recall_batik_mean, record.recall_batik_std),
            pm(record.recall_non_batik_mean, record.recall_non_batik_std),
        ]
        for cell, value in zip(row.cells[1:], values):
            changed += set_cell(cell, value)
    return changed


def refresh_table7(document: Document) -> int:
    """Perbandingan klasik dan frozen-deep pada fold yang sama.

    Urutan baris mengikuti CSV terauditkan agar manuskrip dan sumbernya tidak
    dapat berbeda diam-diam. Karena CSV itu diurutkan menurut macro-F1 CV, dua
    baris handcrafted bertukar tempat setelah fold menjadi group-aware.
    """
    table = find_table(document, {"family", "model", "cv f1", "ext. mcc"})
    source = pd.read_csv(AUDIT_OUT / "table6_external_model_metrics_audited.csv")
    alias = {"SVM (RBF)": "SVM-RBF"}
    source["display"] = source.model.map(lambda name: alias.get(name, name))
    if len(table.rows) - 1 != len(source):
        raise AssertionError("Jumlah baris Table 7 tidak cocok dengan sumber terauditkan")

    posthoc = {
        row.cells[1].text.replace("(post-hoc)", "").strip()
        for row in table.rows[1:]
        if "(post-hoc)" in row.cells[1].text
    }
    families = {
        row.cells[1].text.replace("(post-hoc)", "").strip(): row.cells[0].text.strip()
        for row in table.rows[1:]
    }

    changed = 0
    for row, (_, record) in zip(table.rows[1:], source.iterrows()):
        display = record["display"]
        name = f"{display} (post-hoc)" if display in posthoc else display
        values = [
            families.get(display, row.cells[0].text.strip()),
            name,
            pm(record.cv_f1_macro_mean, record.cv_f1_macro_std),
            f"{fmt(record.external_f1_macro)} "
            f"({fmt(record.external_f1_ci95_low)}-{fmt(record.external_f1_ci95_high)})",
            fmt(record.external_mcc),
            fmt(record.external_recall_batik),
            fmt(record.external_recall_non_batik),
        ]
        for cell, value in zip(row.cells, values):
            changed += set_cell(cell, value)
    return changed


def refresh_table8(document: Document) -> int:
    """Sensitivitas jadwal penyeimbangan kelas."""
    table = find_table(document, {"balancing arm", "training composition", "macro-f1"})
    summary = pd.read_csv(
        RESULTS_DIR / "19_balancing_comparison/arm_summary.csv"
    ).set_index("arm")
    selection = pd.read_csv(
        RESULTS_DIR / "19_balancing_comparison/selection_frequency.csv"
    )
    svm = selection[(selection.kind == "model") & (selection.choice == "SVM (RBF)")]
    counts = dict(zip(svm.arm, svm["count"]))
    of = int(svm["of"].iloc[0])

    labels = {
        "Fold-local augmentation": "augmented_balanced",
        "Class weighting": "class_weighted_originals",
        "Balanced original sampling": "balanced_original_sampling",
    }

    changed = 0
    for row in table.rows[1:]:
        label = row.cells[0].text.strip()
        if label not in labels:
            raise AssertionError(f"Lengan {label!r} tidak dikenal")
        arm = labels[label]
        record = summary.loc[arm]
        values = [
            pm(record.macro_f1_mean, record.macro_f1_std),
            pm(record.balanced_accuracy_mean, record.balanced_accuracy_std),
            fmt(record.mcc_mean),
            f"{int(counts.get(arm, 0))}/{of}",
        ]
        for cell, value in zip(row.cells[2:], values):
            changed += set_cell(cell, value)
    return changed


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if lock_present():
        raise RuntimeError("Manuskrip masih terbuka di Word. Tutup lebih dahulu.")

    document = Document(MANUSCRIPT)
    steps = [
        ("Table 4 diagnostik", refresh_table4),
        ("Table 5 komposisi fold", refresh_table5),
        ("Table 6 hasil 5-fold", refresh_table6),
        ("Table 7 klasik vs deep", refresh_table7),
        ("Table 8 penyeimbangan", refresh_table8),
    ]

    print("=" * 76)
    print("SEGARKAN ANGKA TABEL MANUSKRIP DARI SUMBER TERAUDITKAN")
    print("=" * 76)

    results = []
    total = 0
    for label, step in steps:
        count = step(document)
        results.append((label, count))
        total += count
        print(f"  {label:<28} {count:>3} sel diperbarui")

    if total == 0:
        print("-" * 76)
        print("Tidak ada sel yang berubah. Manuskrip sudah selaras.")
        return

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_number_refresh.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
        print("-" * 76)
        print("Cadangan:", backup.relative_to(PROJECT_DIR))

    document.save(MANUSCRIPT)

    check = Document(MANUSCRIPT)
    red_cells = 0
    for index in range(len(check.tables)):
        for row in check.tables[index].rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip() and run.font.color and run.font.color.rgb == RED:
                            red_cells += 1
                            break
    print("-" * 76)
    print(f"Total sel diperbarui: {total}")
    print(f"Sel bertanda merah setelah penyimpanan: {red_cells}")
    print("Kalimat naratif belum disentuh; itu tugas skrip 27.")


if __name__ == "__main__":
    main()
