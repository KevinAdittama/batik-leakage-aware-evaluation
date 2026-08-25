"""R1.9 - Sisipkan tabel komposisi fold ke manuskrip resmi.

Penempatan yang disetujui: tepat setelah Table 4 (robustness diagnostics) di
bagian Results, sehingga tabel baru menjadi Table 5. Konsekuensinya:

  Table 5 lama (five-fold results)      -> Table 6
  Table 6 lama (classical vs deep)      -> Table 7

Rujukan narasi yang ikut disesuaikan hanya dua, dan keduanya berada pada
paragraf satu-run sehingga tidak memecah content control Mendeley.

Skrip bersifat idempotent: bila tabel R1.9 sudah ada, skrip berhenti tanpa
mengubah apa pun. Seluruh teks yang ditambahkan diberi warna merah sesuai
permintaan Editor butir E.1.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\18_insert_fold_composition_into_manuscript.py
"""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from pipeline_config import PROJECT_DIR

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
TABLE_CSV = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "04_Tabel_Manifest_dan_Hasil"
    / "Audit_Numerik_dan_Eksperimen"
    / "outputs"
    / "R1_9_fold_composition"
    / "fold_composition_table_manuscript.csv"
)
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_R1_9_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)
TABLE_PT = Pt(10)
CAPTION_PT = Pt(9)

MARKER = "Training-fold composition per class"
CAPTION = (
    "Table 5. Training-fold composition per class after the approved two-image "
    "exclusion. Every training split contains exactly 200 instances per class. "
    "Derivatives are training-only descendants of originals in the same training "
    "fold; validation folds contain unaugmented originals only."
)
PROSE = (
    "Training-fold composition is reported per fold and per class in Table 5. "
    "The five validation folds contained 41, 40, 40, 40, and 40 unaugmented "
    "originals, together covering all 201 clean development files exactly once. "
    "Within each training split, batik contributed 109-110 originals and 90-91 "
    "derivatives, whereas non-batik contributed 51-52 originals and 148-149 "
    "derivatives, so both classes reached 200 training instances. Because the "
    "classes differ in the number of available originals, the balancing schedule "
    "drew at most two derivatives per batik original and at most four per "
    "non-batik original, and no original exceeded these per-class limits in any "
    "fold. The resulting asymmetry in derivative density is therefore a property "
    "of the balancing schedule rather than of any individual source."
)

RENUMBER = [
    ("Table 6. Classical and frozen-deep results", "Table 7. Classical and frozen-deep results"),
    ("Table 5. Five-fold results", "Table 6. Five-fold results"),
]
REFERENCE_FIXES = [
    ("retained for descriptive external analysis (Table 5)",
     "retained for descriptive external analysis (Table 6)"),
    ("declined on the separate external collection (Table 6)",
     "declined on the separate external collection (Table 7)"),
]


def body_children(document: Document):
    return list(document.element.body.iterchildren())


def find_table4_element(document: Document):
    """Kembalikan elemen <w:tbl> milik Table 4 berdasarkan caption di atasnya."""
    kids = body_children(document)
    for index, child in enumerate(kids):
        if child.tag.endswith("}p"):
            text = Paragraph(child, document).text.strip()
            if text.startswith("Table 4. Submission robustness diagnostics"):
                for follower in kids[index + 1:]:
                    if follower.tag.endswith("}tbl"):
                        return follower
                    if follower.tag.endswith("}p") and Paragraph(follower, document).text.strip():
                        break
    raise RuntimeError("Table 4 tidak ditemukan di manuskrip")


def styled_run(paragraph, text, *, bold=False, size=TABLE_PT):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.color.rgb = RED
    return run


def build_caption(document: Document) -> Paragraph:
    paragraph = document.add_paragraph(style="Q3 Caption")
    styled_run(paragraph, CAPTION, size=CAPTION_PT)
    return paragraph


def build_prose(document: Document) -> Paragraph:
    paragraph = document.add_paragraph(style="Q3 Body")
    styled_run(paragraph, PROSE, size=Pt(11))
    return paragraph


def build_table(document: Document, frame: pd.DataFrame) -> Table:
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"

    for index, header in enumerate(frame.columns):
        cell = table.rows[0].cells[index]
        cell.text = ""
        styled_run(cell.paragraphs[0], str(header), bold=True)

    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, header in enumerate(frame.columns):
            cells[index].text = ""
            styled_run(cells[index].paragraphs[0], str(row[header]))
    return table


def replace_in_paragraphs(document: Document, pairs) -> list[str]:
    applied = []
    for old, new in pairs:
        for paragraph in document.paragraphs:
            if old not in paragraph.text:
                continue
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    applied.append(f"{old[:52]}... -> {new[:52]}...")
                    break
            break
    return applied


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if not TABLE_CSV.exists():
        raise FileNotFoundError(
            f"{TABLE_CSV} belum ada. Jalankan 17_fold_composition_table.py lebih dahulu."
        )

    document = Document(MANUSCRIPT)
    if any(MARKER in p.text for p in document.paragraphs):
        print("Tabel R1.9 sudah ada di manuskrip. Tidak ada perubahan.")
        return

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_R1_9.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
    print("Cadangan:", backup.relative_to(PROJECT_DIR))

    frame = pd.read_csv(TABLE_CSV)

    # 1. Penomoran mundur lebih dahulu agar tidak bertabrakan.
    renumbered = replace_in_paragraphs(document, RENUMBER)
    fixed_refs = replace_in_paragraphs(document, REFERENCE_FIXES)

    # 2. Bangun elemen baru, lalu pindahkan ke posisi setelah Table 4.
    anchor = find_table4_element(document)
    prose = build_prose(document)
    caption = build_caption(document)
    table = build_table(document, frame)

    anchor.addnext(table._tbl)
    anchor.addnext(caption._p)
    anchor.addnext(prose._p)

    document.save(MANUSCRIPT)

    print("-" * 70)
    print("Penomoran tabel disesuaikan:")
    for item in renumbered:
        print("  ", item)
    print("Rujukan narasi disesuaikan:")
    for item in fixed_refs:
        print("  ", item)
    print("-" * 70)
    print(f"Tabel baru disisipkan sebagai Table 5 dengan {len(frame)} baris data.")
    print("Seluruh teks tambahan berwarna merah (E.1).")


if __name__ == "__main__":
    main()
