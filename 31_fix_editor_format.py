"""E.5 dan E.6 - perbaikan format yang diminta Editor.

E.6 - afiliasi yang berulang
----------------------------
Editor melarang menuliskan afiliasi yang sama berulang kali. Empat penulis
memakai afiliasi yang teksnya identik, tetapi masing-masing diberi nomor sendiri.
Skrip ini menggabungkannya menjadi satu nomor yang dirujuk bersama, lalu
menomori ulang afiliasi yang tersisa dan menyesuaikan superskrip pada setiap nama.

Pemetaannya diturunkan dari teks afiliasinya sendiri, bukan ditulis harfiah,
sehingga tetap benar bila daftar afiliasi berubah kemudian.

E.5 - ukuran huruf judul bagian
-------------------------------
Editor meminta seluruh kalimat 11 pt. Satu judul bagian memakai 12 pt sedangkan
judul lain 11 pt. Ketidakseragaman itu diseragamkan ke 11 pt.

Label (a) dan (b) di dalam gambar ditangani skrip 29, bukan di sini.

Perubahan ditandai merah sesuai butir E.1, tetapi hanya pada bagian yang benar-
benar berubah. Nomor superskrip yang kebetulan tidak bergeser dibiarkan hitam.

Skrip bersifat idempotent dan menolak jalan bila manuskrip terbuka di Word.

Jalankan dari folder proyek:
    .\\env\\Scripts\\python.exe -B .\\31_fix_editor_format.py
"""

from __future__ import annotations

import shutil
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor

from pipeline_config import PROJECT_DIR

MANUSCRIPT = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "01_Dokumen_Siap_Submit"
    / (
        "Leakage-Aware Evaluation Reveals Acquisition Bias and External "
        "Degradation in Binary Batik Recognition.docx"
    )
)
BACKUP_DIR = (
    PROJECT_DIR
    / "IJIES_REVISI_FINAL"
    / "99_Arsip_Pendukung"
    / f"Versi_Sebelum_Format_Editor_{date.today():%Y%m%d}"
)

RED = RGBColor(0xFF, 0x00, 0x00)
BODY_POINT = 11.0
HEADER_SCAN = 16  # paragraf awal yang memuat nama dan afiliasi


def lock_present() -> bool:
    return any(MANUSCRIPT.parent.glob("~$*"))


def affiliation_paragraphs(document: Document) -> list:
    """Paragraf afiliasi: superskrip angka di depan, lalu teks lembaganya."""
    found = []
    for paragraph in document.paragraphs[:HEADER_SCAN]:
        runs = [run for run in paragraph.runs if run.text.strip()]
        if len(runs) < 2:
            continue
        first = runs[0]
        if not first.font.superscript or not first.text.strip().isdigit():
            continue
        body = "".join(run.text for run in runs[1:]).strip()
        if "Department" not in body and "Faculty" not in body:
            continue
        found.append((paragraph, first, body))
    return found


def consolidate_affiliations(document: Document) -> tuple[int, dict]:
    entries = affiliation_paragraphs(document)
    if not entries:
        raise AssertionError("Paragraf afiliasi tidak ditemukan")

    order: list[str] = []
    for _, _, body in entries:
        if body not in order:
            order.append(body)
    new_number = {body: index for index, body in enumerate(order, 1)}
    mapping = {
        marker.text.strip(): new_number[body] for _, marker, body in entries
    }

    seen: set[str] = set()
    changed = 0
    for paragraph, marker, body in entries:
        if body in seen:
            paragraph._p.getparent().remove(paragraph._p)
            changed += 1
            continue
        seen.add(body)
        target = str(new_number[body])
        if marker.text.strip() != target:
            marker.text = target
            marker.font.color.rgb = RED
            changed += 1
    return changed, mapping


def retarget_author_markers(document: Document, mapping: dict) -> int:
    """Sesuaikan superskrip pada nama penulis mengikuti penomoran baru."""
    changed = 0
    affiliation_ids = {
        id(paragraph) for paragraph, _, _ in affiliation_paragraphs(document)
    }
    for paragraph in document.paragraphs[:HEADER_SCAN]:
        if id(paragraph) in affiliation_ids:
            continue
        text = paragraph.text.strip()
        if "Department" in text or "@" in text or not text:
            continue
        for run in paragraph.runs:
            token = run.text.strip()
            if not run.font.superscript or not token.isdigit():
                continue
            if token not in mapping:
                raise AssertionError(f"Superskrip {token!r} tanpa afiliasi padanan")
            target = str(mapping[token])
            if token != target:
                run.text = run.text.replace(token, target)
                run.font.color.rgb = RED
                changed += 1
    return changed


# Singkatan yang dilarang Editor pada nama dan afiliasi, beserta bentuk penuhnya.
# Sengaja dibuat sangat sempit: hanya singkatan yang jelas, bukan koreksi ejaan
# nama tempat. Kesalahan ejaan pada alamat mitra bukan urusan skrip ini.
ABBREVIATIONS = {"Jln": "Jalan", "Jl.": "Jalan", "Univ.": "University"}


def expand_abbreviations(document: Document) -> list[str]:
    """Editor melarang singkatan pada nama penulis dan afiliasi."""
    touched = []
    for paragraph in document.paragraphs[:HEADER_SCAN]:
        for run in paragraph.runs:
            for short, full in ABBREVIATIONS.items():
                if short not in run.text:
                    continue
                run.text = run.text.replace(short, full)
                run.font.color.rgb = RED
                touched.append(f"{short} -> {full}")
    return touched


def unify_heading_size(document: Document) -> list[str]:
    """Turunkan judul bagian yang lebih besar dari badan teks ke 11 pt."""
    touched = []
    for paragraph in document.paragraphs:
        oversized = [
            run for run in paragraph.runs
            if run.text.strip() and run.font.size and run.font.size.pt > BODY_POINT
        ]
        if not oversized:
            continue
        for run in oversized:
            run.font.size = Pt(BODY_POINT)
        touched.append(paragraph.text.strip()[:48])
    return touched


def main() -> None:
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    if lock_present():
        raise RuntimeError("Manuskrip masih terbuka di Word. Tutup lebih dahulu.")

    document = Document(MANUSCRIPT)

    print("=" * 76)
    print("E.5 DAN E.6 - PERBAIKAN FORMAT YANG DIMINTA EDITOR")
    print("=" * 76)

    before = affiliation_paragraphs(document)
    unique = len({body for _, _, body in before})
    print(f"  afiliasi tertulis: {len(before)} | teks unik: {unique}")

    mapping_changes, mapping = consolidate_affiliations(document)
    marker_changes = retarget_author_markers(document, mapping)
    abbreviations = expand_abbreviations(document)
    headings = unify_heading_size(document)

    total = mapping_changes + marker_changes + len(abbreviations) + len(headings)
    if total == 0:
        print("-" * 76)
        print("Tidak ada perubahan. Format sudah sesuai.")
        return

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / "Manuscript_before_editor_format.docx"
    if not backup.exists():
        shutil.copy2(MANUSCRIPT, backup)
        print("-" * 76)
        print("Cadangan:", backup.relative_to(PROJECT_DIR))

    document.save(MANUSCRIPT)

    check = Document(MANUSCRIPT)
    after = affiliation_paragraphs(check)
    print("-" * 76)
    print("pemetaan nomor afiliasi lama -> baru:",
          ", ".join(f"{old}->{new}" for old, new in sorted(mapping.items())))
    print(f"  afiliasi setelah digabung: {len(after)}")
    for _, marker, body in after:
        print(f"    {marker.text.strip()}  {body[:66]}")
    print(f"  superskrip nama disesuaikan: {marker_changes}")
    if abbreviations:
        print(f"  singkatan pada afiliasi dibuka: {abbreviations}")
    if headings:
        print(f"  judul diseragamkan ke {BODY_POINT:.0f} pt: {headings}")

    duplicates = len(after) - len({body for _, _, body in after})
    if duplicates:
        raise AssertionError(f"Masih ada {duplicates} afiliasi berulang")
    print("-" * 76)
    print("Tidak ada afiliasi yang ditulis berulang.")
    print("Tata letak halaman judul perlu diperiksa visual di Word.")


if __name__ == "__main__":
    main()
